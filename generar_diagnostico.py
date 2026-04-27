from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Helpers ───────────────────────────────────────────────────────────────────
def heading(doc, text, size, color):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*color)
    return p

def bullet(doc, text, bold_part=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_part:
        rb = p.add_run(bold_part)
        rb.bold = True
    p.add_run(text)

def code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
    p.paragraph_format.left_indent = Inches(0.4)

def table_header_row(tbl, cols):
    row = tbl.rows[0].cells
    for i, col in enumerate(cols):
        row[i].text = col
        for run in row[i].paragraphs[0].runs:
            run.bold = True

def add_table_row(tbl, values):
    row = tbl.add_row().cells
    for i, v in enumerate(values):
        row[i].text = v

AZUL   = (0x1E, 0x40, 0xAF)
AZUL2  = (0x1D, 0x4E, 0x89)
VERDE  = (0x14, 0x53, 0x14)
NARANJA= (0xB4, 0x53, 0x09)
GRIS   = (0x99, 0x99, 0x99)

# ── PORTADA ───────────────────────────────────────────────────────────────────
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("WMS-PAME")
r.bold = True; r.font.size = Pt(30); r.font.color.rgb = RGBColor(*AZUL)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = s.add_run("Diagnóstico de Cambios — Módulo Packing & Picking")
r2.font.size = Pt(14); r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
f.add_run(f"Fecha: {datetime.date.today().strftime('%d de %B de %Y')}   |   Ambiente: Railway Production")
doc.add_paragraph()

# ── 1. RESUMEN EJECUTIVO ──────────────────────────────────────────────────────
heading(doc, "1.  Resumen Ejecutivo", 15, AZUL)
doc.add_paragraph(
    "Durante esta sesión de desarrollo se corrigieron errores críticos en el módulo de Packing "
    "del WMS-PAME y se implementaron mejoras en la experiencia del operario empacador. "
    "Se realizaron 5 commits a la rama main con deploy automático en Railway Production."
)

# ── 2. CONTEXTO — SESIÓN ANTERIOR ────────────────────────────────────────────
doc.add_paragraph()
heading(doc, "2.  Cambios de Sesión Anterior (Contexto)", 15, AZUL)
doc.add_paragraph("Los siguientes cambios ya estaban en producción como base de esta sesión:")

previos = [
    ("Botón «Reiniciar conteo» en Picking:",
     " Permite resetear el conteo de unidades y paquetes a 0 sin cancelar la tarea."),
    ("Botón «Reiniciar conteo» en Packing:",
     " Cada tarjeta de pedido revierte todos sus ítems y el estado (incluyendo VERIFICADO → EN_PROCESO)."),
    ("Fix: empaques_escaneados no se reiniciaba:",
     " Se corrigió reseteando también el contador grande de paquetes en backend y frontend."),
    ("Fix: HUD packing «Error de conexión» tras reiniciar:",
     " Se reemplazó empCargarTareas() por empCerrarHUD() para limpiar el estado global del HUD."),
    ("Fix: tarjeta EMPAQUE SUGERIDO redundante:",
     " Se ocultaba para productos sin empaque (factor ≤ 1)."),
    ("Fix: botón Confirmar deshabilitado al recargar con tarea completa:",
     " Se pre-evalúa el estado inicial en renderTarea()."),
    ("HUD Packing — línea de unidades:",
     " Se agregó #emp-hud-und mostrando «X / Y und» para productos paquete."),
    ("Logs Railway:",
     " Se configuró logging.basicConfig(stream=sys.stdout) y gunicorn con --access-logfile - --error-logfile -."),
]
for bold_part, rest in previos:
    bullet(doc, rest, bold_part=bold_part)

# ── 3. CAMBIOS DE ESTA SESIÓN ─────────────────────────────────────────────────
doc.add_paragraph()
heading(doc, "3.  Cambios de Esta Sesión", 15, AZUL)

# 3.1
doc.add_paragraph()
heading(doc, "3.1  Fix — Error real visible al fallar un escaneo de packing", 12, AZUL2)
doc.add_paragraph(
    "Problema: el bloque catch del HUD de empaque mostraba siempre «Error de conexión» sin importar "
    "el error real del servidor, imposibilitando el diagnóstico en campo."
)
doc.add_paragraph("Solución aplicada:")
bullet(doc, "app/static/pwa/app.js:", bold_part="")
code(doc, 'catch (e) { empFlash("rojo", e.message && e.message !== "401" ? e.message : "Error de conexión"); }')
bullet(doc, "app/services/mobile_service.py:", bold_part="")
code(doc, "item.cantidad_real = (item.cantidad_real or 0) + cantidad  # defensivo contra NULL en DB")
doc.add_paragraph(
    "El operario y el administrador ahora ven el mensaje exacto del servidor "
    "(ej. «Producto X no pertenece a este pedido»)."
)

# 3.2
doc.add_paragraph()
heading(doc, "3.2  Fix — Packing acepta todos los códigos del producto", 12, AZUL2)
doc.add_paragraph(
    "Problema: el endpoint /api/mobile/escanear solo buscaba por codigo, codigo_siesa y codigo_barras. "
    "Si el operario escaneaba el código de barras del empaque (codigo_barras_empaque), "
    "el backend respondía «Producto no encontrado»."
)
doc.add_paragraph(
    "Solución: el bloque PACKING de MobileService.procesar_escaneo ahora itera los ítems "
    "de la tarea usando _codigos_validos(), que incluye los cuatro tipos de código del producto. "
    "Además, si el código escaneado corresponde al empaque y el frontend no aplicó el factor, "
    "el backend lo aplica automáticamente."
)
code(doc, "for it in tarea.items:")
code(doc, "    if p and codigo_limpio in MobileService._codigos_validos(p):")
code(doc, "        producto = p; item = it; break")

# 3.3
doc.add_paragraph()
heading(doc, "3.3  Feat — HUD Packing muestra paquetes en grande (igual que Picking)", 12, AZUL2)
doc.add_paragraph(
    "Problema: el HUD de empaque no mostraba el conteo en paquetes para productos con "
    "factor_conversion > 1. El contador siempre mostraba unidades sueltas."
)
doc.add_paragraph(
    "Causa raíz: EMP_EMPAQUES se cargaba desde /api/empaques/descomponer (tabla productos_empaques), "
    "pero muchos productos tienen el factor solo en productos.factor_conversion "
    "(parametrizado directamente en Siesa), por lo que el endpoint devolvía factor = 1."
)
doc.add_paragraph("Cambios:")
bullet(doc,
    "ItemPacking.to_dict() ahora incluye factor_conversion y unidad_empaque del producto.",
    bold_part="Backend — app/models/packing.py:  ")
bullet(doc,
    "empIniciarHUD() ya no llama a /api/empaques/descomponer; popula EMP_EMPAQUES "
    "directamente desde item.factor_conversion.",
    bold_part="Frontend — app/static/pwa/app.js:  ")
doc.add_paragraph("Resultado visual (ejemplo factor = 36, pedido de 72 und):")
bullet(doc, "Número grande cyan: 1  (paquetes escaneados)")
bullet(doc, "Línea debajo: de 2 CAJA")
bullet(doc, "Segunda línea gris: 36 / 72 und")

# 3.4
doc.add_paragraph()
heading(doc, "3.4  Feat — Packing solo puede contar lo que salió del Picking", 12, AZUL2)
doc.add_paragraph(
    "Requerimiento: el packing es una segunda verificación — debe contar exactamente lo que "
    "el picker recogió, ni más ni menos."
)

tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(tbl, ["Nivel", "Regla", "Implementación"])
rows = [
    ("Origen del dato",
     "cantidad_esperada del packing = cantidad_recogida del picking (no la cantidad ordenada en Siesa)",
     "PackingService.crear_desde_picking usa tp.cantidad_recogida"),
    ("Límite superior",
     "No se puede escanear más de lo que el picker entregó",
     "Exceso check en MobileService.procesar_escaneo y PackingService.escanear_item"),
    ("Límite inferior",
     "No se puede confirmar si algún ítem no llegó a la cantidad esperada",
     "confirmar_packing verifica items_sin_verificar antes de marcar VERIFICADO"),
    ("Garantía",
     "El packing solo cierra cuando cantidad_real == cantidad_esperada para cada ítem",
     "Combinación de los tres niveles anteriores"),
]
for r in rows:
    add_table_row(tbl, r)

# ── 4. COMMITS ────────────────────────────────────────────────────────────────
doc.add_paragraph()
heading(doc, "4.  Commits Realizados — rama main → Railway", 15, AZUL)

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = "Table Grid"
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(tbl2, ["SHA", "Mensaje", "Archivos"])
commits = [
    ("ce322a1", "fix: mostrar error real en packing scan y defensivo contra NULL",
     "mobile_service.py, app.js"),
    ("1b6c773", "fix: packing scan busca por todos los codigos del producto",
     "mobile_service.py"),
    ("4ac401a", "feat: packing HUD muestra paquetes usando factor_conversion directo",
     "models/packing.py, app.js"),
    ("e9c82e6", "feat: packing limita escaneo a cantidad_esperada del picking",
     "mobile_service.py"),
    ("d9bf646", "fix: bloquear exceso en endpoint admin — garantizar packing == picking",
     "packing_service.py"),
]
for sha, msg, files in commits:
    row = tbl2.add_row().cells
    row[0].text = sha
    row[0].paragraphs[0].runs[0].font.name = "Courier New"
    row[0].paragraphs[0].runs[0].font.size = Pt(9)
    row[1].text = msg
    row[2].text = files

# ── 5. PENDIENTES ─────────────────────────────────────────────────────────────
doc.add_paragraph()
heading(doc, "5.  Pendientes", 15, AZUL)

# 5.1 Guía
doc.add_paragraph()
heading(doc, "5.1  Guía de Usuario — Módulo Packing", 12, NARANJA)
doc.add_paragraph(
    "Se debe redactar la guía operativa para el empacador. Temas a cubrir:"
)
guia_items = [
    "Flujo completo: abrir una tarea de packing desde la lista de pedidos",
    "Cómo escanear productos con láser Bluetooth y con cámara",
    "Interpretación del HUD: número grande (paquetes) vs. unidades debajo",
    "Qué hacer ante errores comunes: exceso, producto incorrecto, error de conexión",
    "Botón «Reiniciar conteo» — cuándo usarlo y consecuencias",
    "Paso de declarar bultos físicos al cerrar la caja",
    "Confirmación visual de remisión generada en Siesa",
    "Rol del jefe de almacén vs. empacador en el flujo de excepciones",
]
for item in guia_items:
    bullet(doc, item)

# 5.2 Facturación
doc.add_paragraph()
heading(doc, "5.2  Facturación del Pedido — Integración Siesa (PENDIENTE DE VALIDACIÓN)", 12, NARANJA)
doc.add_paragraph(
    "El trigger de facturación electrónica existe en el código pero debe ser validado "
    "y documentado completamente antes de habilitarse en producción:"
)
fact_items = [
    "Validar el trigger FacturaPedido (ID 238925) en Connekta: QA → Producción",
    "Verificar payload: tipo_docto_pedido, consec_docto_pedido, items (item_id_siesa, unidad_medida, cantidades)",
    "Manejar respuestas de Siesa: número de remisión generado, errores de stock y crédito",
    "Retry automático: estado DESPACHADO con siesa_triggered=False debe reintentarse sin duplicar remisión",
    "Mostrar al empacador el número de remisión generado tras el cierre",
    "Alertas al administrador cuando Siesa rechaza la facturación (crédito agotado, producto sin parametrizar, etc.)",
    "Pruebas end-to-end en ambiente QA con pedidos reales antes de activar en producción",
    "Documentar casos borde: pedidos con múltiples items, productos con lote, short-picks confirmados",
]
for item in fact_items:
    bullet(doc, item)

# ── 6. ARCHIVOS MODIFICADOS ───────────────────────────────────────────────────
doc.add_paragraph()
heading(doc, "6.  Archivos Modificados Esta Sesión", 15, AZUL)

tbl3 = doc.add_table(rows=1, cols=2)
tbl3.style = "Table Grid"
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(tbl3, ["Archivo", "Descripción del cambio"])
archivos = [
    ("app/services/mobile_service.py",
     "Bloque PACKING: búsqueda por _codigos_validos, detección de empaque, exceso check, defensivo NULL"),
    ("app/services/packing_service.py",
     "escanear_item: exceso check + corrección de lógica verificado"),
    ("app/models/packing.py",
     "ItemPacking.to_dict(): agrega factor_conversion y unidad_empaque del producto"),
    ("app/static/pwa/app.js",
     "empIniciarHUD: EMP_EMPAQUES desde item.factor_conversion; catch blocks con e.message"),
]
for archivo, cambio in archivos:
    row = tbl3.add_row().cells
    row[0].text = archivo
    row[0].paragraphs[0].runs[0].font.name = "Courier New"
    row[0].paragraphs[0].runs[0].font.size = Pt(9)
    row[1].text = cambio

# ── PIE ───────────────────────────────────────────────────────────────────────
doc.add_paragraph()
pie = doc.add_paragraph()
pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
pr = pie.add_run("WMS-PAME  ·  Diagnóstico generado con Claude Code  ·  2026")
pr.font.size = Pt(9)
pr.font.color.rgb = RGBColor(*GRIS)

# ── GUARDAR ───────────────────────────────────────────────────────────────────
ruta = r"C:\Users\SSJUAN03\Desktop\WMS-PAME\Diagnostico_Packing_2026-04-23.docx"
doc.save(ruta)
print("Documento generado:", ruta)
