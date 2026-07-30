"""
Genera DIAGNOSTICO_PICKING_PARCIAL.docx
Ejecutar: python generar_diagnostico_word.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Estilos globales ──────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(10.5)

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colores ───────────────────────────────────────────────────────────────────
AZUL_SIESA   = RGBColor(0x1A, 0x73, 0xE8)   # azul Connekta
ROJO         = RGBColor(0xC0, 0x39, 0x2B)
VERDE        = RGBColor(0x27, 0xAE, 0x60)
NARANJA      = RGBColor(0xE6, 0x7E, 0x22)
GRIS_FONDO   = RGBColor(0xF2, 0xF2, 0xF2)
AZUL_CABEZA  = RGBColor(0x1F, 0x49, 0x7D)
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)
NEGRO        = RGBColor(0x00, 0x00, 0x00)
AMARILLO_BG  = RGBColor(0xFF, 0xF3, 0xCD)


def set_cell_bg(cell, color: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, border_color='1F497D', size=6):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(size))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), border_color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = AZUL_CABEZA
    # Línea inferior
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = AZUL_SIESA
    return p


def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = AZUL_CABEZA
    return p


def body(text, bold_parts=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if bold_parts:
        parts = text.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.bold = (i % 2 == 1)
    else:
        p.add_run(text)
    return p


def code_block(lines):
    """Bloque de código con fondo gris."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        # Fondo gris claro
        pPr  = p._p.get_or_add_pPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F2F2F2')
        pPr.append(shd)


def alerta(texto, tipo='warning'):
    """Caja de alerta: warning=naranja, error=rojo, ok=verde, info=azul."""
    colores = {
        'warning': (NARANJA,  'FFF3CD'),
        'error':   (ROJO,     'FADBD8'),
        'ok':      (VERDE,    'D5F5E3'),
        'info':    (AZUL_SIESA,'D6EAF8'),
    }
    color_txt, hex_bg = colores.get(tipo, colores['info'])
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_bg)
    tcPr.append(shd)
    p    = cell.paragraphs[0]
    run  = p.add_run(texto)
    run.font.color.rgb = color_txt
    run.bold = True
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(0.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def tabla_datos(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Cabecera
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        set_cell_bg(cell, AZUL_CABEZA)
        p    = cell.paragraphs[0]
        run  = p.add_run(h)
        run.bold = True
        run.font.color.rgb = BLANCO
        run.font.size = Pt(9.5)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Filas
    for i, row in enumerate(rows):
        bg = GRIS_FONDO if i % 2 == 0 else BLANCO
        for j, val in enumerate(row):
            cell = t.cell(i + 1, j)
            set_cell_bg(cell, bg)
            p    = cell.paragraphs[0]
            run  = p.add_run(str(val))
            run.font.size = Pt(9.5)
    # Anchos
    if col_widths:
        for i, row_cells in enumerate(t.rows):
            for j, cell in enumerate(row_cells.cells):
                cell.width = Inches(col_widths[j])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def checklist(items):
    """items: list of (done: bool, text: str)"""
    for done, text in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        icono = '✅' if done else '☐'
        run = p.add_run(f'{icono}  {text}')
        run.font.size = Pt(10)
        if done:
            run.font.color.rgb = VERDE
        else:
            run.font.color.rgb = NEGRO


# ═══════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(30)
run = p_title.add_run('DIAGNÓSTICO TÉCNICO')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = AZUL_CABEZA

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p_sub.add_run('Despacho Parcial — Conector 142945 ignora f470_cant_base')
run2.bold = True
run2.font.size = Pt(14)
run2.font.color.rgb = ROJO

doc.add_paragraph()
p_meta = doc.add_paragraph()
p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p_meta.add_run(
    f'Proyecto: WMS-PAME — Papelería Medellín\n'
    f'Fecha: {datetime.date.today().strftime("%d/%m/%Y")}\n'
    f'Ambiente: QA → integradorqa.siesacloud.com\n'
    f'Archivos: app/services/connekta_gateway.py · despacho_parcial_service.py'
)
run3.font.size = Pt(10)
run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  1. CONTEXTO DEL PROBLEMA
# ═══════════════════════════════════════════════════════════════════════════════
h1('1. Contexto del problema')

body('El WMS envía despachos parciales vía conector **142945** (API_v1_Ventas_Comercial_RemisionPedido). '
     'El payload incluye la cantidad real picada por el operario y el rowid de la línea del pedido:', bold_parts=True)

code_block([
    '{',
    '  "f470_cant_base": 3.0,',
    '  "f470_rowid_movto": 470406',
    '}',
])

doc.add_paragraph()
tabla_datos(
    ['Campo', 'Valor'],
    [
        ['Resultado esperado', 'SIESA graba 3 unidades en T470'],
        ['Resultado real', 'SIESA graba 4 unidades (cantidad comprometida completa)'],
    ],
    col_widths=[2.5, 4.0]
)

# ═══════════════════════════════════════════════════════════════════════════════
#  2. INVESTIGACIÓN — LO QUE SE DESCARTÓ
# ═══════════════════════════════════════════════════════════════════════════════
h1('2. Investigación realizada — lo que se descartó')

h2('2.1 Perfil de usuario Connekta (m.agudelo)')
items_desc = [
    '✅  Usuario m.agudelo → perfil Administrador confirmado en Administrador de Seguridad SIESA',
    '✅  Checkbox "No almacenar cantidades parciales al aprob y comp" → UNCHECKED en perfil Administrador',
    '✅  Conclusión: la configuración de seguridad es correcta. No es el problema.',
]
for t in items_desc:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(t)
    run.font.size = Pt(10)

h2('2.2 Código del WMS')
items_cod = [
    '✅  f470_rowid_movto se envía correctamente (valor real: 470406 = f431_rowid de T431)',
    '✅  f470_cant_base se envía con el valor real picado (3.0)',
    '✅  Commit 6028daa verificó que rowid_map retorna valores reales',
    '✅  Conclusión: el código del WMS es correcto. No hay código que tocar.',
]
for t in items_cod:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.add_run(t).font.size = Pt(10)

h2('2.3 Backorder del cliente')
body('Proceso de estandarización activo con Siesa (2026-05-20). No es la causa del problema de cantidad. Son issues separados.')

# ═══════════════════════════════════════════════════════════════════════════════
#  3. CAUSA RAÍZ
# ═══════════════════════════════════════════════════════════════════════════════
h1('3. Causa raíz — diagnóstico definitivo')

h2('3.1 Arquitectura interna de SIESA para despachos parciales')
body('La tabla T405 (t405_cm_compromisos) tiene dos campos de cantidad:')

tabla_datos(
    ['Campo', 'Valor en el escenario', 'Descripción'],
    [
        ['f405_cant_por_remisionar_base', '4.0', 'Cantidad comprometida por SIESA al aprobar el pedido'],
        ['f405_cant_picking_base',        'NULL', 'Cantidad real picada por el WMS — NO fue escrita'],
    ],
    col_widths=[2.8, 1.5, 3.2]
)

body('El conector 142945 tiene lógica interna que:', bold_parts=False)
pasos = [
    '1. Recibe f470_rowid_movto (= f431_rowid) para encontrar la fila en T431.',
    '2. Desde T431 navega a T405.',
    '3. Lee f405_cant_picking_base → si es NULL → usa f405_cant_por_remisionar_base = 4.',
    '4. Ignora f470_cant_base del payload cuando f405_cant_picking_base es NULL.',
]
for t in pasos:
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    p.add_run(t).font.size = Pt(10)

h2('3.2 El botón "Parciales" en la UI de SIESA')
body('El flujo manual que hace un usuario en SIESA replica exactamente lo que necesitamos:')
pasos_parciales = [
    '1. Abre pantalla Remisiones desde Pedidos.',
    '2. Presiona botón "Parciales".',
    '3. SIESA escribe f405_cant_picking_base = 3 en T405 para ese compromiso.',
    '4. Crea la remisión → 142945 lee f405_cant_picking_base = 3 → graba 3 en T470.',
]
for t in pasos_parciales:
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    p.add_run(t).font.size = Pt(10)

alerta('⚠️  El WMS necesita replicar el paso 3 de forma programática ANTES de llamar 142945.', 'warning')

# ═══════════════════════════════════════════════════════════════════════════════
#  4. SOLUCIÓN EN CONNEKTA QA
# ═══════════════════════════════════════════════════════════════════════════════
h1('4. Solución implementada en Connekta QA')

h2('4.1 Lo que NO se puede hacer — Conectores dinámicos')
body('En Apis Dinámicas Siesa → Conectores solo se pueden importar conectores del catálogo oficial '
     'de SIESA. No se pueden crear conectores nuevos desde cero. No existe en el catálogo un conector '
     'que actualice f405_cant_picking_base.')

h2('4.2 Lo que SÍ se hizo — Consulta dinámica (UPDATE vía Módulo Conectividad)')
body('En Apis Dinámicas Siesa → Consultas se creó la siguiente consulta:')

tabla_datos(
    ['Campo', 'Valor'],
    [
        ['ID',          '7811'],
        ['Nombre',      'papeleriamedellin_WMS_set_picking_parcial'],
        ['Conexión',    'Módulo conectividad (SQL Server local)'],
        ['Versión',     '3.0'],
        ['Estado',      '✅ Activo'],
    ],
    col_widths=[2.0, 4.5]
)

h3('SQL configurado:')
code_block([
    'UPDATE t405_cm_compromisos',
    'SET    f405_cant_picking_base = @cantidad',
    'WHERE  f405_id_cia = 1',
    '  AND  f405_rowid_pv_movto = @f431_rowid',
    '  AND  f405_cant_por_remisionar_base >= @cantidad',
    '  AND  @cantidad > 0',
])
doc.add_paragraph()

h3('Guards de seguridad del SQL:')
guards = [
    'f405_id_cia = 1 → solo empresa PAME, nunca cruza datos',
    'f405_rowid_pv_movto = @f431_rowid → fila exacta del compromiso, no masivo',
    'f405_cant_por_remisionar_base >= @cantidad → imposible despachar más de lo comprometido',
    '@cantidad > 0 → nunca borra el picking con un 0 accidental',
]
for t in guards:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.add_run(t).font.size = Pt(10)

# ═══════════════════════════════════════════════════════════════════════════════
#  5. BLOQUEO ACTUAL
# ═══════════════════════════════════════════════════════════════════════════════
h1('5. Bloqueo actual — Permiso pendiente de Siesa Soporte')

h2('5.1 Por qué falla ahora')
body('La consulta 7811 existe y está activa. Pero cuando el WMS la llame, el agente local del '
     'Módulo de Conectividad se conecta a SQL Server con un usuario de servicio distinto a m.agudelo. '
     'Ese usuario NO tiene permiso UPDATE sobre dbo.t405_cm_compromisos.')

code_block([
    'WMS → GET /api/connekta/v3/ejecutarconsulta (ID: 7811)',
    '  → Connekta Cloud',
    '    → Módulo Conectividad Siesa SQL Server (agente local)',
    '      → SQL Server: SUnoEE_Papeleriamed_Imple',
    '        → UPDATE t405_cm_compromisos  ← ❌ ERROR: permiso denegado',
])
doc.add_paragraph()

alerta('🔴  BLOQUEADO: Se requiere GRANT UPDATE sobre dbo.t405_cm_compromisos para el usuario del Módulo de Conectividad.', 'error')

# ═══════════════════════════════════════════════════════════════════════════════
#  6. MENSAJE PARA SIESA SOPORTE
# ═══════════════════════════════════════════════════════════════════════════════
h1('6. Mensaje para Siesa Soporte (listo para enviar)')

body('Asunto: Habilitar permiso UPDATE en t405_cm_compromisos para Módulo de Conectividad — QA y Producción', bold_parts=False)

code_block([
    'Hola equipo,',
    '',
    'En el ambiente QA (integradorqa.siesacloud.com) creamos la consulta dinámica',
    'ID 7811 — papeleriamedellin_WMS_set_picking_parcial. Esta consulta ejecuta un',
    'UPDATE sobre dbo.t405_cm_compromisos (campo f405_cant_picking_base) a través',
    'del Módulo de Conectividad Siesa SQL Server.',
    '',
    'Lo que necesitamos:',
    'Que el usuario de BD del Módulo de Conectividad tenga permiso de escritura:',
    '',
    '  GRANT UPDATE ON dbo.t405_cm_compromisos',
    '      TO [<usuario_servicio_modulo_conectividad>];',
    '',
    'Contexto técnico:',
    'El WMS necesita escribir f405_cant_picking_base antes de ejecutar el conector',
    '142945 (RemisionPedido), para que SIESA respete la cantidad real picada por el',
    'operario y no la cantidad comprometida completa. Esto replica el comportamiento',
    'del botón "Parciales" de Remisiones desde Pedidos.',
    '',
    'Preguntas adicionales:',
    '1. ¿Cuál es el nombre del usuario SQL Server del Módulo de Conectividad?',
    '2. ¿Las consultas dinámicas UPDATE/DML son soportadas por ejecutarconsulta,',
    '   o se requiere un endpoint diferente?',
    '3. ¿Este permiso aplica también en producción o se gestiona por separado?',
])
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  7. FLUJO COMPLETO PEDIDOS PARCIALES
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('7. Flujo completo — Pedidos Parciales WMS-PAME')

h2('7.1 ¿Quién lo dispara?')
body('Solo usuarios con **rol de gestión** (_es_gestion()). Es un flujo exclusivo del administrador, '
     'completamente separado del flujo del operario (cerrar_caja).', bold_parts=True)

h2('7.2 Flujo principal')

# Tabla de flujo paso a paso
tabla_datos(
    ['Paso', 'Acción', 'Detalle técnico', 'Estado'],
    [
        ['0',  'Admin consulta compromisos (opcional)',
         'GET /api/despacho_parcial/<id>/compromisos\nAPI_v2_Ventas_Pedidos_Compromisos → muestra cantidades comprometidas en Siesa',
         '✅ OK'],
        ['1',  'Admin ejecuta despacho',
         'POST /api/despacho_parcial/<id>/despachar\nBody: {"cantidades": {"PAPELSP9218": 3.0}}',
         '✅ OK'],
        ['2',  'Obtener cabecera del pedido',
         'API_v2_Ventas_Pedidos → tercero, moneda, cond. pago, f430_rowid',
         '✅ OK'],
        ['3',  'Obtener rowid_map',
         'API_v2_Ventas_Pedidos_Compromisos filtrado por f430_rowid\n→ {PAPELSP9218: 470406}',
         '✅ OK'],
        ['4',  'Construir items del despacho',
         '_build_items(): prioriza cantidades del body (WMS truth)\nFallback: _build_items_from_compromisos() si WMS sin cantidades',
         '✅ OK'],
        ['5',  'Idempotencia — ¿ya se procesó?',
         'rm_tipo+rm_consec en BD → saltar 142945\nCompromisos vacíos → buscar FE existente',
         '✅ OK'],
        ['6',  '⚠️  SET picking en T405',
         'Consulta 7811: UPDATE t405 SET f405_cant_picking_base=3\nWHERE f431_rowid=470406\nPrerequisito para que 142945 respete f470_cant_base',
         '🔴 PENDIENTE PERMISO'],
        ['7',  'POST 142945 → RemisionPedido',
         'f470_cant_base=3, f470_rowid_movto=470406\nResponse: "Importacion exitosa" (sin consecutivo)',
         '⚠️ Funciona pero graba 4'],
        ['8',  'Auto-detect consecutivo RM',
         'Consulta 7479 (papeleriamedellin_WMS_Remision_DesdePedido)\nSQL: JOIN t350→t470→t431→t430 → {tipo:"RM", consec:1458}',
         '✅ OK'],
        ['9',  'Guardar RM en BD',
         'tarea.rm_tipo="RM", tarea.rm_consec=1458 → commit independiente\nGarantiza retry sin re-crear RM si falla el paso 10',
         '✅ OK'],
        ['10', 'Guard anti-duplicado FE',
         'get_factura_desde_pedido() → ¿ya existe FE? → retornar sin 142943',
         '✅ OK'],
        ['11', 'POST 142943 → FacturaDesdeRemision',
         'Input: tipo_rm="RM", consec_rm=1458\n→ FE creada en Siesa → tarea.estado=DESPACHADO\n→ tarea.siesa_triggered=True → commit BD',
         '✅ OK'],
    ],
    col_widths=[0.4, 2.0, 3.5, 1.1]
)

h2('7.3 Carriles de recuperación')

tabla_datos(
    ['Carril', 'Cuándo usarlo', 'Endpoint', 'Qué hace'],
    [
        ['A — Facturar remisión existente',
         'RM existe en Siesa y en BD,\npero la FE no se creó',
         'POST /api/despacho_parcial/<id>/facturar-remision\n(sin body)',
         'Lee rm_tipo/rm_consec de BD → dispara 142943 directamente'],
        ['B — Facturar RM manual',
         'RM existe en Siesa, pero el\nconsecutivo NO está en BD',
         'POST /api/despacho_parcial/<id>/facturar-rm-manual\nBody: {"tipo_rm":"RM","consec_rm":1458}',
         'Admin busca nro. en Siesa → guarda en BD → dispara 142943'],
    ],
    col_widths=[1.8, 1.8, 2.2, 2.2]
)

h2('7.4 El bug activo — por qué T470 graba 4 en vez de 3')

code_block([
    'Paso 6 (set_picking_parcial) NO existe todavía  ← falta permiso Siesa Soporte',
    '    ↓',
    'f405_cant_picking_base = NULL en T405',
    '    ↓',
    '142945 lee f405_cant_por_remisionar_base = 4',
    '    ↓',
    'Ignora f470_cant_base = 3 del payload',
    '    ↓',
    'T470 graba 4  ← BUG',
])
doc.add_paragraph()

h2('7.5 Resumen de conectores y consultas del flujo parcial')

tabla_datos(
    ['Paso', 'Conector / Consulta', 'Tipo', 'Acción'],
    [
        ['Compromisos', 'API_v2_Ventas_Pedidos_Compromisos',          'GET estándar',       'Lee compromisos y f431_rowid'],
        ['Cabecera',   'API_v2_Ventas_Pedidos',                       'GET estándar',       'Lee datos del cliente y f430_rowid'],
        ['⚠️ Paso 6',  'Consulta 7811 — WMS_set_picking_parcial',     'GET dinámico (DML)', 'UPDATE T405 cant_picking (pendiente permiso)'],
        ['RM',         '142945 — RemisionPedido',                     'POST conector',      'Crea remisión de salida'],
        ['Auto-detect','Consulta 7479 — WMS_Remision_DesdePedido',    'GET dinámico',       'Encuentra consecutivo RM en BD Siesa'],
        ['FE',         '142943 — FacturaDesdeRemision',               'POST conector',      'Convierte RM → Factura Electrónica'],
    ],
    col_widths=[0.7, 2.8, 1.4, 2.1]
)

# ═══════════════════════════════════════════════════════════════════════════════
#  8. CÓDIGO WMS LISTO
# ═══════════════════════════════════════════════════════════════════════════════
h1('8. Código WMS listo para integrar (pendiente permiso)')

h2('8.1 connekta_gateway.py — nuevo método set_picking_parcial()')

code_block([
    'def set_picking_parcial(self, f431_rowid: int, cantidad: float) -> dict:',
    '    """',
    '    Consulta dinámica ID 7811 — papeleriamedellin_WMS_set_picking_parcial',
    '    PREREQUISITO OBLIGATORIO antes de llamar trigger_despacho() (142945).',
    '    Escribe f405_cant_picking_base en T405 para que 142945 respete',
    '    f470_cant_base en lugar de f405_cant_por_remisionar_base.',
    '    Equivale al botón "Parciales" de la UI de Remisiones desde Pedidos.',
    '    """',
    '    if self.modo_simulacion:',
    '        logger.info("[CONNEKTA SIM] set_picking_parcial: rowid=%s cant=%s",',
    '                    f431_rowid, cantidad)',
    '        return {"simulado": True}',
    '',
    '    return self._get(',
    '        "papeleriamedellin_WMS_set_picking_parcial",',
    '        params_extra={"parametros": f"f431_rowid={f431_rowid}|cantidad={cantidad}"},',
    '        url=self.url_get_dinamico,',
    '    )',
])
doc.add_paragraph()

h2('8.2 despacho_parcial_service.py — insertar antes del bloque de 142945')

code_block([
    '# Paso 0 — escribir cant_picking en T405 (prerequisito para 142945)',
    '# Sin esto, 142945 usa f405_cant_por_remisionar_base (cant. comprometida)',
    '# en lugar de respetar el f470_cant_base que enviamos.',
    'for ref, rowid in rowid_map.items():',
    '    cant = float(cantidades.get(ref, 0))',
    '    if cant > 0 and rowid:',
    '        try:',
    '            connekta.set_picking_parcial(f431_rowid=rowid, cantidad=cant)',
    '            logger.info("[DESPACHO_PARCIAL] T405 picking seteado: ref=%s cant=%s",',
    '                        ref, cant)',
    '        except Exception as e:',
    '            # No bloquear — loguear y continuar (fallback a cant. comprometida)',
    '            logger.warning("[DESPACHO_PARCIAL] set_picking_parcial falló: %s", e)',
])
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  9. INVENTARIO DE CONSULTAS Y CONECTORES
# ═══════════════════════════════════════════════════════════════════════════════
h1('9. Inventario Connekta QA')

h2('9.1 Consultas dinámicas activas')
tabla_datos(
    ['ID', 'Nombre', 'Uso'],
    [
        ['6719', 'papeleriamedellin_WMS_Picking_Pedidos_NB1',  'Pedidos en estado picking para bodega NB1'],
        ['7461', 'papeleriamedellin_test_remisiones_temp',      'Testing temporal — evaluar eliminación'],
        ['7479', 'papeleriamedellin_WMS_Remision_DesdePedido',  'Auto-detect consecutivo RM post-142945'],
        ['7798', 'papeleriamedellin_compromisos_wms',           'Mapeo {f431_rowid: f405_rowid} para pedidos'],
        ['7799', 'papeleriamedellin_pame_descubrir_tablas',     'Introspección de tablas — uso diagnóstico'],
        ['7811', 'papeleriamedellin_WMS_set_picking_parcial',   '🔴 UPDATE T405 — NUEVO, pendiente permiso'],
    ],
    col_widths=[0.6, 3.2, 2.8]
)

h2('9.2 Conectores dinámicos activos')
body('⚠️  En conectores solo se puede IMPORTAR del catálogo SIESA. No se pueden crear nuevos.')
tabla_datos(
    ['ID', 'Nombre', 'Módulo principal', 'Sub-módulo'],
    [
        ['238920', 'CLASIFICACION DE ITEMS',                  '01_Maestro',   '3_Comercial'],
        ['238925', 'FACTURA_DESDE_PEDIDO',                    '03_Comercial', '2_Ventas'],
        ['239216', 'API_v1_Ventas_Comercial_RemisionPedido',  '12_Connekta',  '6_ConectoresEstandar'],
        ['244114', 'ACTUALIZACION_BACKORDER_PAME_V1',         '01_Maestro',   '3_Comercial'],
    ],
    col_widths=[0.8, 3.0, 1.5, 1.8]
)

# ═══════════════════════════════════════════════════════════════════════════════
#  10. CHECKLIST
# ═══════════════════════════════════════════════════════════════════════════════
h1('10. Checklist de próximos pasos')

checklist([
    (True,  'Consulta 7811 creada y activa en Connekta QA'),
    (True,  'SQL validado con guards de seguridad correctos'),
    (True,  'Código WMS preparado (método set_picking_parcial + integración en despacho_parcial_service)'),
    (False, 'BLOQUEADO: Siesa Soporte habilita GRANT UPDATE en t405 para usuario Módulo Conectividad'),
    (False, 'Confirmar con Soporte que ejecutarconsulta soporta DML (UPDATE), no solo SELECT'),
    (False, 'Una vez habilitado: integrar set_picking_parcial en connekta_gateway.py'),
    (False, 'Test manual en QA: despacho parcial 3 uds → verificar T470 = 3 (no 4)'),
    (False, 'Verificar log: "[DESPACHO_PARCIAL] T405 picking seteado"'),
    (False, 'Si test OK: commit + deploy Railway QA'),
    (False, 'Replicar consulta en producción (servicios.siesacloud.com) con mismo SQL'),
    (False, 'Replicar permiso GRANT en producción'),
    (False, 'Test end-to-end producción con 1 pedido de prueba'),
])

# ═══════════════════════════════════════════════════════════════════════════════
#  11. REFERENCIAS DE CÓDIGO
# ═══════════════════════════════════════════════════════════════════════════════
h1('11. Referencias de código')

tabla_datos(
    ['Archivo', 'Método / Función', 'Propósito'],
    [
        ['connekta_gateway.py',          'get_pedido_rowid_map()',         'Genera {ref: f431_rowid} por pedido'],
        ['connekta_gateway.py',          'get_compromisos_pedido()',       'Obtiene f431_rowid desde API_v2_Ventas_Pedidos_Compromisos'],
        ['connekta_gateway.py',          'trigger_despacho()',             'POST 142945 — usa f470_rowid_movto y f470_cant_base'],
        ['connekta_gateway.py',          'get_remision_desde_pedido()',    'Auto-detect consecutivo RM (consulta 7479)'],
        ['connekta_gateway.py',          'get_compromisos_t405()',         'Consulta dinámica compromisos_wms'],
        ['despacho_parcial_service.py',  'despachar_parcial()',            'Orquestador principal del flujo parcial'],
        ['despacho_parcial_service.py',  '_build_items()',                 'Construye items con rowid_movto desde WMS'],
        ['despacho_parcial_service.py',  'facturar_remision_existente()',  'Carril A — RM existe en BD, genera FE'],
        ['despacho_parcial_service.py',  'facturar_rm_con_consec()',       'Carril B — consecutivo manual de emergencia'],
        ['routes/despacho_parcial.py',   '/despachar (POST)',              'Endpoint principal — rol gestión requerido'],
        ['routes/despacho_parcial.py',   '/facturar-remision (POST)',      'Endpoint carril A'],
        ['routes/despacho_parcial.py',   '/facturar-rm-manual (POST)',     'Endpoint carril B'],
    ],
    col_widths=[2.2, 2.5, 2.8]
)

# ── Pie de página ─────────────────────────────────────────────────────────────
doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_footer = p_footer.add_run(
    f'WMS-PAME — Papelería Medellín  ·  Generado {datetime.date.today().strftime("%d/%m/%Y")}  '
    '·  Siguiente acción: ticket a Siesa Soporte (Sección 6)'
)
run_footer.font.size = Pt(8)
run_footer.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Guardar ───────────────────────────────────────────────────────────────────
output = 'DIAGNOSTICO_PICKING_PARCIAL.docx'
doc.save(output)
print(f'[OK]  Documento generado: {output}')
