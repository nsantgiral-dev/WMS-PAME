"""
Genera MODULO_DEVOLUCIONES_RECONCILIACION.docx
Ejecutar: python generar_doc_devoluciones.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Estilos globales ──────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(10.5)

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colores ───────────────────────────────────────────────────────────────────
AZUL_SIESA   = RGBColor(0x1A, 0x73, 0xE8)
ROJO         = RGBColor(0xC0, 0x39, 0x2B)
VERDE        = RGBColor(0x27, 0xAE, 0x60)
NARANJA      = RGBColor(0xE6, 0x7E, 0x22)
GRIS_FONDO   = RGBColor(0xF2, 0xF2, 0xF2)
AZUL_CABEZA  = RGBColor(0x1F, 0x49, 0x7D)
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)
NEGRO        = RGBColor(0x00, 0x00, 0x00)


def set_cell_bg(cell, color: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = AZUL_CABEZA
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = AZUL_SIESA
    return p


def body(text, bold_parts=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if bold_parts:
        parts = text.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.bold = (i % 2 == 1)
    else:
        p.add_run(text)
    return p


def paso(numero, titulo, detalle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    run_num = p.add_run(f'{numero}.  ')
    run_num.bold = True
    run_num.font.color.rgb = AZUL_SIESA
    run_tit = p.add_run(f'{titulo} — ')
    run_tit.bold = True
    p.add_run(detalle)


def alerta(texto, tipo='info'):
    colores = {
        'warning': (NARANJA,  'FFF3CD'),
        'error':   (ROJO,     'FADBD8'),
        'ok':      (VERDE,    'D5F5E3'),
        'info':    (AZUL_SIESA, 'D6EAF8'),
    }
    color_txt, hex_bg = colores.get(tipo, colores['info'])
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_bg)
    tcPr.append(shd)
    p    = cell.paragraphs[0]
    run  = p.add_run(texto)
    run.font.color.rgb = color_txt
    run.bold = True
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(0.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def tabla_datos(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        set_cell_bg(cell, AZUL_CABEZA)
        p    = cell.paragraphs[0]
        run  = p.add_run(h)
        run.bold = True
        run.font.color.rgb = BLANCO
        run.font.size = Pt(9.5)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        bg = GRIS_FONDO if i % 2 == 0 else BLANCO
        for j, val in enumerate(row):
            cell = t.cell(i + 1, j)
            set_cell_bg(cell, bg)
            p    = cell.paragraphs[0]
            run  = p.add_run(str(val))
            run.font.size = Pt(9.5)
    if col_widths:
        for row_cells in t.rows:
            for j, cell in enumerate(row_cells.cells):
                cell.width = Inches(col_widths[j])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def bullet(text, bold_parts=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_parts:
        parts = text.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.bold = (i % 2 == 1)
    else:
        p.add_run(text)


# ═══════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(30)
run = p_title.add_run('MÓDULO DE DEVOLUCIONES POR RECONCILIACIÓN')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = AZUL_CABEZA

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p_sub.add_run('Logística Inversa — corrección física de inventario contra Siesa')
run2.bold = True
run2.font.size = Pt(13)
run2.font.color.rgb = AZUL_SIESA

doc.add_paragraph()
p_meta = doc.add_paragraph()
p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p_meta.add_run(
    f'Proyecto: WMS-PAME — Papelería Medellín\n'
    f'Fecha: {datetime.date.today().strftime("%d/%m/%Y")}\n'
    f'Rol involucrado: Recepcionista (pestaña "Devoluciones")'
)
run3.font.size = Pt(10)
run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
h1('1. ¿Qué es este módulo?')
body(
    'Es el mecanismo del WMS para corregir diferencias de inventario cuando '
    '**Siesa (el sistema contable/ERP) tiene registrada más cantidad de un '
    'producto de la que el WMS tiene ubicada físicamente en bodega**.'
)
body(
    'Regla de negocio central: **este módulo no toca dinero, IVA ni contabilidad.** '
    'Siesa ya procesó esa cantidad — el WMS solo se encarga de encontrarle una '
    'ubicación física real en el estante y dejar el inventario del sistema '
    'igualado a la realidad.'
)
alerta(
    'En corto: Siesa dice "tienes más de este producto de lo que registras" → '
    'el recepcionista confirma dónde está físicamente esa mercancía → el WMS '
    'actualiza su inventario.',
    tipo='info'
)

# ═══════════════════════════════════════════════════════════════════════════════
h1('2. ¿Cuándo se genera una tarea de devolución?')
body(
    'Un administrador ejecuta una reconciliación de inventario, que descarga el '
    'stock real desde Siesa y lo compara contra lo que el WMS tiene registrado, '
    'producto por producto. El resultado se clasifica así:'
)
tabla_datos(
    headers=['Clasificación', 'Significado', '¿Genera tarea aquí?'],
    rows=[
        ['SIESA_MAYOR', 'Siesa tiene más cantidad que el WMS', 'SÍ — es este módulo'],
        ['WMS_MAYOR',   'El WMS tiene más cantidad que Siesa', 'No — se revisa aparte'],
        ['SOLO_WMS',    'Producto que el WMS tiene pero no llegó en la descarga de Siesa', 'No — se revisa aparte'],
    ],
    col_widths=[1.6, 3.4, 1.8]
)
body(
    'Para evitar falsas alarmas, el sistema solo genera tareas automáticamente si '
    'el WMS ya tiene cargado al menos el 20% del catálogo de Siesa — si la '
    'cobertura es menor, el resultado queda solo como informativo para revisión manual.'
)

# ═══════════════════════════════════════════════════════════════════════════════
h1('3. Flujo paso a paso')
paso('1', 'Reconciliación',
     'Un administrador dispara la reconciliación de inventario contra Siesa.')
paso('2', 'Detección',
     'El sistema identifica los productos donde Siesa reporta más cantidad que el WMS.')
paso('3', 'Creación de tarea',
     'Por cada producto con diferencia se crea una "Tarea de Devolución" con la '
     'cantidad exacta de la diferencia — sin duplicar si ya existe una tarea activa '
     'para ese mismo producto.')
paso('4', 'Visualización',
     'El recepcionista ve la tarea en la pestaña "Devoluciones" de su panel, con la '
     'diferencia resaltada (ej. "+140").')
paso('5', 'Ubicación física',
     'El recepcionista localiza físicamente la mercancía sobrante, toca la tarea y '
     'elige (o escanea) la ubicación de bodega donde la va a dejar. También puede '
     'marcarla como "averiada" si el producto no está en condiciones de venta.')
paso('6', 'Actualización de inventario',
     'El sistema suma la cantidad a esa ubicación, registra el movimiento para '
     'trazabilidad, y cierra la tarea como Completada.')
paso('7', 'Caso especial — averiado',
     'Si se marca como averiado, además de mover el producto a una zona de '
     'cuarentena, el sistema notifica a Siesa el traslado a la bodega de averías '
     '(única excepción donde este módulo sí interactúa con Siesa).')

# ═══════════════════════════════════════════════════════════════════════════════
h1('4. Caso real observado hoy')
body(
    f'Al {datetime.date.today().strftime("%d/%m/%Y")}, la pestaña de Devoluciones del '
    'recepcionista mostraba 2 tareas pendientes de ubicar:'
)
tabla_datos(
    headers=['Producto', 'Referencia', 'Diferencia (Siesa vs WMS)'],
    rows=[
        ['MARCADOR PELIKAN 420 NEGRO', 'PAPELSP6741', '+140 unidades'],
        ['TABLA DE APOYO FYC PLÁSTICA VERDE', 'PAPELSP9964', '+15 unidades'],
    ],
    col_widths=[3.0, 1.8, 2.2]
)
body(
    'Ambos productos están registrados en Siesa con más cantidad que la que el WMS '
    'tiene ubicada — hasta que el recepcionista confirme dónde están físicamente, '
    'el WMS seguirá subestimando el stock disponible de esas dos referencias.'
)

# ═══════════════════════════════════════════════════════════════════════════════
h1('5. Estado operativo actual')
alerta(
    'Este proceso es 100% manual hoy: tanto la reconciliación (que detecta las '
    'diferencias) como la carga de stock inicial desde Siesa se disparan por '
    'botón — no hay una tarea programada que las ejecute automáticamente cada '
    'cierto tiempo.',
    tipo='warning'
)
body(
    'Implicación práctica: si nadie ejecuta la reconciliación con cierta '
    'frecuencia, estas diferencias pueden acumularse sin que el recepcionista '
    'se entere, y el stock disponible que ve el WMS puede quedar por debajo del '
    'real durante días.'
)

# ── Pie de página ─────────────────────────────────────────────────────────────
doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_footer = p_footer.add_run(
    f'WMS-PAME — Papelería Medellín  ·  Generado {datetime.date.today().strftime("%d/%m/%Y")}'
)
run_footer.font.size = Pt(8)
run_footer.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Guardar ───────────────────────────────────────────────────────────────────
output = 'MODULO_DEVOLUCIONES_RECONCILIACION.docx'
doc.save(output)
print(f'[OK]  Documento generado: {output}')
