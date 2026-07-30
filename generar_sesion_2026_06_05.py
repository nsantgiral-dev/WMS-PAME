"""Genera el documento de sesión 2026-06-05 — WMS-PAME."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Estilos globales ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def p(text=''):
    return doc.add_paragraph(text)

def code(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    para.paragraph_format.left_indent = Inches(0.4)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    para._p.get_or_add_pPr().append(shd)
    return para

def tabla(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t

def bullet(text, bold_prefix=None):
    para = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = para.add_run(bold_prefix)
        run.bold = True
        para.add_run(text)
    else:
        para.add_run(text)
    return para

# ═══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading('WMS-PAME — Sesión 2026-06-05', 0)
p('Solución del Bloqueante B1: Conector RIT 174646')
p(f'Generado: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}')
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. RESUMEN EJECUTIVO
# ═══════════════════════════════════════════════════════════════════════════════
h1('1. Resumen Ejecutivo')
p('En esta sesión se resolvió el bloqueante B1 del flujo de traslados ST: '
  'el conector Connekta 174646 (RequisicionesParaTransferir) generaba registros '
  'del tipo 441 con 2557 caracteres, pero Siesa Enterprise v5 exige exactamente '
  '2614 caracteres. Después de múltiples intentos de corrección de valores, se '
  'diagnosticó que el problema era estructural en el serializador de Connekta, '
  'y se construyó un conector custom en UnoEE Generic Transfer (WMS_PAME_RIT_v6) '
  'que incluye el layout v6 completo con los 34 campos requeridos.')
p()
p('Al cierre de la sesión el conector custom está creado y activo en Siesa QA '
  '(ID 245835). Pendiente: actualizar variables Railway y ejecutar prueba e2e.')

# ═══════════════════════════════════════════════════════════════════════════════
# 2. CONTEXTO PREVIO
# ═══════════════════════════════════════════════════════════════════════════════
h1('2. Contexto — Estado al Inicio de Sesión')
p('El flujo ST (traslados entre bodegas) estaba parcialmente operativo:')
tabla(
    ['Paso', 'Estado', 'Detalle'],
    [
        ['Crear solicitud', '✅ OK', 'Tienda crea en BORRADOR → ENVIADA'],
        ['Admin aprueba', '✅ OK', 'Transición a EN_PICKING'],
        ['174646 RIT', '❌ BLOQUEADO', '2557 chars vs 2614 requeridos'],
        ['Picking NS1', '✅ OK', 'Pickers ven tareas ST'],
        ['Packing NS1', '✅ OK', 'Cerrar caja funciona'],
        ['DLQ → STS (174930)', '⏳ BLOQUEADO', 'Sin consec_rit, cae a fallback 173076'],
        ['173076 fallback', '❌ FALLA', 'Siesa exige "Movimiento de Seriales"'],
        ['ETS recepción tienda', '⏳ NO PROBADO', 'Bloqueado por pasos anteriores'],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. DIAGNÓSTICO DEL PROBLEMA
# ═══════════════════════════════════════════════════════════════════════════════
h1('3. Diagnóstico del Problema 174646')

h2('3.1 Síntoma')
p('Siesa retorna HTTP 400 con:')
code('El tamaño del registro no corresponde al exigido.\nTamaño del registro = 2557. Tamaño registro exigido = 2614.')

h2('3.2 Intentos de corrección de valores (todos fallidos)')
tabla(
    ['Commit', 'Cambio aplicado', 'Resultado'],
    [
        ['14d90e6', 'None → "" para 3 campos trailing', '2557 — sin cambio'],
        ['3dcd912', '"" → " " (espacio) para trailing', '2557 — sin cambio'],
        ['4e94628', 'Mover RIT a confirmar_picking con ubicacion real ("NS1")', '2557 — sin cambio'],
    ]
)
p()
p('Conclusión: el serializador de Connekta 174646 NO procesa los 3 campos trailing '
  'independientemente del valor enviado. Es un defecto del conector estándar.')

h2('3.3 Análisis del Layout v6 — Hallazgos clave')
p('Se obtuvo el layout completo del registro tipo 441 versión 6 desde UnoEE Generic Transfer:')

tabla(
    ['Campo', 'Posición', 'Ancho', 'Obligatorio', 'Estado en 174646'],
    [
        ['f441_id_ubicacion_sal', '2555', '10', 'No', '❌ No serializado'],
        ['f441_id_proy_etapa', '2565', '30', 'Dep', '❌ No serializado'],
        ['f441_id_rubro_pof', '2595', '20', 'No', '❌ No serializado'],
        ['f441_id_moneda_sug', '2615', '3', 'Dep', '(v6 nuevo — sin 174646)'],
    ]
)
p()
p('Los 3 primeros campos combinan 60 chars (10+30+20). El registro v5 termina en '
  'posición 2614 = f441_id_rubro_pof final. Connekta 174646 no serializa ninguno '
  'de estos campos, produciendo siempre 2557 chars.')

p('Otros hallazgos del layout:')
bullet(' f441_desc_varible (pos 515, ancho 2000) — no está en el JSON de WMS; '
       'Connekta lo rellena con 2000 espacios automáticamente. Explica el salto de pos 514 a 2515.', 'IMPORTANTE: ')
bullet(' f441_id_concepto = 607 (fijo) — también manejado internamente por Connekta.', 'FIJO: ')
bullet(' Errores de "campos obligatorios en blanco" en posiciones 194-214 (f441_cant_2), '
       '2515-2534 (f441_id_un_movto) y 2535-2554 (f441_precio_unitario): '
       'Connekta serializa 0 como espacios para decimales opcionales.', 'SECUNDARIO: ')

# ═══════════════════════════════════════════════════════════════════════════════
# 4. CAMBIO ARQUITECTÓNICO: RIT EN CONFIRMAR_PICKING
# ═══════════════════════════════════════════════════════════════════════════════
h1('4. Cambio Arquitectónico — RIT Movida a confirmar_picking')

h2('4.1 Decisión')
p('El consultor Siesa (SIESA-014) y el layout v6 confirmaron que f441_id_ubicacion_sal '
  'debe contener la ubicación real del bin escaneado durante el picking. '
  'La RIT se estaba creando en aprobar_solicitud() con ubicacion_codigo=None '
  '(sin picking aún). Se movió al momento correcto: confirmar_picking_traslado().')

h2('4.2 Cambio en traslado_service.py (commit 4e94628)')
tabla(
    ['Función', 'Antes', 'Después'],
    [
        ['aprobar_solicitud()', 'Crea RIT con ubicacion=None', 'Solo crea tasks de picking'],
        ['confirmar_picking_traslado()', 'No llama a 174646', 'Llama RIT con ubicaciones reales del bin'],
    ]
)
p()
p('Lógica de ubicación en confirmar_picking_traslado():')
code("ubicacion_codigo = _ubicaciones_rit.get(item.producto_id) or s.bodega_origen_siesa\n"
     "# Usa ubicacion real del WMS; fallback al código de bodega origen")

# ═══════════════════════════════════════════════════════════════════════════════
# 5. SOLUCIÓN: CONECTOR CUSTOM WMS_PAME_RIT_v6
# ═══════════════════════════════════════════════════════════════════════════════
h1('5. Solución — Conector Custom WMS_PAME_RIT_v6')

h2('5.1 Decisión de diseño')
p('En lugar de esperar soporte de Connekta, se construyó un conector custom en '
  'UnoEE Generic Transfer basado en el layout v6 completo del plano '
  'API_v1_Inventarios_Comercial_RequisicionesParaTransferir. '
  'El v6 incluye los 34 campos con posiciones correctas, incluyendo los 3 campos '
  'que Connekta 174646 omitía.')

h2('5.2 Datos del conector creado')
tabla(
    ['Parámetro', 'Valor'],
    [
        ['ID Connekta (idDocumento)', '245835'],
        ['Nombre (nombreDocumento)', 'WMS_PAME_RIT_v6'],
        ['Plano base', 'API_v1_Inventarios_Comercial_RequisicionesParaTransferir'],
        ['Módulo', '12_Connekta → 6_ConectoresEstandar'],
        ['ERP', 'UnoEE'],
        ['Estado', 'Activo (punto azul en portal Conectores Dinámicos)'],
        ['Creado', '05/06/2026 11:24:00'],
    ]
)

h2('5.3 Configuración de secciones')
tabla(
    ['Sección', 'Versión', 'Observación'],
    [
        ['Inicial', '1', 'Sin cambios vs 174646'],
        ['Documentos', '4', 'Sin cambios vs 174646 — f440_id_clase_docto=75 fijo'],
        ['Movimientos', '6', 'Layout completo 34 campos — F_VERSION-REG=06'],
        ['Final', '1', 'Sin cambios vs 174646'],
    ]
)

h2('5.4 Campos fijos configurados en Movimientos v6')
tabla(
    ['Campo', 'Valor fijo', 'Descripción'],
    [
        ['F_TIPO-REG', '0441', 'Tipo de registro Movimientos'],
        ['F_SUBTIPO-REG', '00', 'Subtipo fijo'],
        ['F_VERSION-REG', '06', 'Versión layout — CAMBIADO de 05 a 06'],
        ['f441_id_concepto', '607', 'Transferencias'],
        ['f_campo', '(vacío)', 'Pos 228, 2 chars — campos futuros'],
        ['f441_desc_varible', '(vacío)', 'Pos 515, 2000 chars — descripción'],
        ['f441_id_moneda_sug', '(vacío)', 'Pos 2615, 3 chars — v6 nuevo, Dep'],
    ]
)

h2('5.5 Campos que ahora se serializan correctamente')
tabla(
    ['Campo', 'Posición', 'Ancho', 'Valor enviado desde WMS'],
    [
        ['f441_id_ubicacion_sal', '2555', '10', 'ubicacion real del picking (fallback: bodega origen)'],
        ['f441_id_proy_etapa', '2565', '30', 'None (PAME no usa proyectos)'],
        ['f441_id_rubro_pof', '2595', '20', 'None (PAME no usa presupuesto POF)'],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 6. CAMBIOS DE CÓDIGO
# ═══════════════════════════════════════════════════════════════════════════════
h1('6. Cambios de Código')

tabla(
    ['Commit', 'Archivo', 'Cambio'],
    [
        ['4e94628', 'traslado_service.py', 'RIT movida de aprobar_solicitud a confirmar_picking_traslado con ubicaciones reales'],
        ['f30483f', 'connekta_gateway.py', 'Agregar self.nombre_conector_req_traslado desde env CONNEKTA_NOMBRE_REQ_TRASLADO; usar en _post()'],
    ]
)

p()
h2('Cambio en connekta_gateway.py')
code("# Línea ~53 — nuevas variables de inicialización:\n"
     "self.conector_requisicion_traslado = os.getenv('CONNEKTA_CONECTOR_REQ_TRASLADO', '174646')\n"
     "self.nombre_conector_req_traslado  = os.getenv('CONNEKTA_NOMBRE_REQ_TRASLADO',\n"
     "                                                'API_v1_Inventarios_Comercial_RequisicionesParaTransferir')\n\n"
     "# Línea ~1967 — llamada al conector:\n"
     "return self._post(self.conector_requisicion_traslado,\n"
     "                  self.nombre_conector_req_traslado, payload)")

# ═══════════════════════════════════════════════════════════════════════════════
# 7. VARIABLES RAILWAY
# ═══════════════════════════════════════════════════════════════════════════════
h1('7. Variables Railway — Configuración Requerida')

p('⚠️ AL CIERRE DE SESIÓN ESTAS VARIABLES AÚN NO ESTABAN ACTUALIZADAS EN RAILWAY. '
  'La prueba ST-20260605-ECFB usó el conector viejo (174646) porque el deploy '
  'ejecutó con los valores anteriores.')

tabla(
    ['Variable', 'Valor anterior', 'Valor correcto', 'Estado'],
    [
        ['CONNEKTA_CONECTOR_REQ_TRASLADO', '174646', '245835', '⚠️ PENDIENTE actualizar'],
        ['CONNEKTA_NOMBRE_REQ_TRASLADO', '(no existía)', 'WMS_PAME_RIT_v6', '⚠️ PENDIENTE crear'],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 8. AGENDA PRÓXIMA SESIÓN
# ═══════════════════════════════════════════════════════════════════════════════
h1('8. Agenda Próxima Sesión')

h2('CRÍTICO — Primera acción')
bullet('Actualizar Railway: CONNEKTA_CONECTOR_REQ_TRASLADO=245835 + crear CONNEKTA_NOMBRE_REQ_TRASLADO=WMS_PAME_RIT_v6')
bullet('Verificar en log: "conector=245835 nombre=WMS_PAME_RIT_v6"')
bullet('Crear solicitud ST nueva → aprobar → picker completa → confirmar picking → revisar log RIT')

h2('Si RIT retorna HTTP 200 con consecutivo')
bullet('Verificar que s.siesa_requisicion_consec se guarda en BD')
bullet('Ejecutar flujo completo: packing → cerrar caja → DLQ → 174930 STS → 173079 ETS')
bullet('Confirmar que 174930 no exige seriales (a diferencia de 173076)')

h2('Si RIT falla con nuevo error de Siesa v6')
bullet('El error dirá el nuevo tamaño requerido o campos inválidos')
bullet('Si dice "f_version 6 no reconocida": el conector UnoEE debe usar F_VERSION-REG=05 en lugar de 06')
bullet('  → Cambiar en UnoEE: Movimientos v6, campo F_VERSION-REG, FIJO: 05', '  Opción: ')

h2('Deuda técnica pendiente (post e2e ST)')
tabla(
    ['ID', 'Descripción', 'Prioridad'],
    [
        ['DT-E2E-ST', 'Prueba e2e completa flujo ST en producción', 'CRÍTICA'],
        ['DT-PACK-01', 'Eliminar _cerrar_packing_pedido_legacy()', 'Media'],
        ['DT-PACK-03', 'Revisar/eliminar confirmar_packing_traslado()', 'Media'],
        ['DT-BUGS', 'Limpiar 5 jobs DESPACHO_F470 FALLIDO históricos', 'Baja'],
        ['DT-INV-46', 'Investigar discrepancia 46 uds PAPELSP9218', 'Baja'],
        ['DT-ZONA', '_buscar_ubicacion_optima filtra por zona vs tipo_zona', 'Media'],
        ['Migración 007', 'Agregar compromisos_ok, consec_sts, sts_ok a SiesaJob', 'Pendiente auth'],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 9. REFERENCIA TÉCNICA — LAYOUT 441 v6
# ═══════════════════════════════════════════════════════════════════════════════
h1('9. Referencia — Layout Completo Registro 441 v6')

tabla(
    ['#', 'Campo', 'Pos inicio', 'Ancho', 'Tipo', 'Obligatorio'],
    [
        ['1', 'F_NUMERO-REG', '1', '7', 'Cglobal', 'Si'],
        ['2', 'F_TIPO-REG', '8', '4', 'Entero (=0441)', 'Si'],
        ['3', 'F_SUBTIPO-REG', '12', '2', 'Entero (=00)', 'Si'],
        ['4', 'F_VERSION-REG', '14', '2', 'Entero (=06)', 'Si'],
        ['5', 'F_CIA', '16', '3', 'Entero', 'Si'],
        ['6', 'f441_id_co', '19', '3', 'Alfanumérico', 'Si'],
        ['7', 'f441_id_tipo_docto', '22', '3', 'Alfanumérico', 'Si'],
        ['8', 'f441_consec_docto', '25', '8', 'Entero', 'Si'],
        ['9', 'f441_nro_registro', '33', '10', 'Entero', 'Si'],
        ['10', 'f441_id_item', '43', '7', 'Entero', 'Dep'],
        ['11', 'f441_referencia_item', '50', '50', 'Alfanumérico', 'Dep'],
        ['12', 'f441_codigo_barras', '100', '20', 'Alfanumérico', 'Dep'],
        ['13', 'f441_id_ext1_detalle', '120', '20', 'Alfanumérico', 'Dep'],
        ['14', 'f441_id_ext2_detalle', '140', '20', 'Alfanumérico', 'Dep'],
        ['15', 'f441_id_bodega', '160', '5', 'Alfanumérico', 'Si'],
        ['16', 'f441_id_concepto', '165', '3', 'Entero (=607)', 'Si'],
        ['17', 'f441_id_motivo', '168', '2', 'Alfanumérico', 'Si'],
        ['18', 'f441_id_unidad_medida', '170', '4', 'Alfanumérico', 'Si'],
        ['19', 'f441_cant_base', '174', '20', 'Decimal 15.4', 'Si'],
        ['20', 'f441_cant_2', '194', '20', 'Decimal 15.4', 'Dep'],
        ['21', 'f441_fecha_entrega', '214', '8', 'Alfanumérico AAAAMMDD', 'Si'],
        ['22', 'f441_num_dias_entrega', '222', '3', 'Entero', 'No'],
        ['23', 'f441_id_co_movto', '225', '3', 'Alfanumérico', 'Si'],
        ['24', 'f_campo', '228', '2', 'Alfanumérico (vacío)', 'No'],
        ['25', 'f441_id_ccosto_movto', '230', '15', 'Alfanumérico', 'Dep'],
        ['26', 'f441_id_proyecto', '245', '15', 'Alfanumérico', 'No'],
        ['27', 'f441_notas', '260', '255', 'Alfanumérico', 'No'],
        ['28', 'f441_desc_varible', '515', '2000', 'Alfanumérico', 'No'],
        ['29', 'f441_id_un_movto', '2515', '20', 'Alfanumérico', 'Si'],
        ['30', 'f441_precio_unitario', '2535', '20', 'Decimal 15.4', 'Dep'],
        ['31', 'f441_id_ubicacion_sal', '2555', '10', 'Alfanumérico', 'No'],
        ['32', 'f441_id_proy_etapa', '2565', '30', 'Alfanumérico', 'Dep'],
        ['33', 'f441_id_rubro_pof', '2595', '20', 'Alfanumérico', 'No'],
        ['34', 'f441_id_moneda_sug', '2615', '3', 'Alfanumérico', 'Dep'],
    ]
)

p()
p('Total record v6: 2617 chars (posición 2615 + ancho 3 - 1 = 2617).')
p('Total record v5: 2614 chars (hasta f441_id_rubro_pof, posición 2595 + ancho 20 - 1 = 2614).')

# ── Guardar ───────────────────────────────────────────────────────────────────
ruta = r'C:\Users\SSJUAN03\Desktop\WMS-PAME\SESION_WMS_PAME_2026-06-05.docx'
doc.save(ruta)
print(f'Documento generado: {ruta}')
