"""Genera SESION_2026-06-01.docx con el resumen de la sesión WMS-PAME."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

doc = Document()

# ── Estilos base ──────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def body(text):
    return doc.add_paragraph(text)

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.add_run(text)
    return p

def tabla(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # Header
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1d4ed8')
        tcPr.append(shd)
    # Filas
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        for ci, val in enumerate(row):
            tr.cells[ci].text = str(val)
    return t

def separador():
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
titulo = doc.add_paragraph()
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = titulo.add_run('WMS-PAME — Resumen de Sesión')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run('Fecha: 2026-06-01  ·  Módulo: Rutas / Muelle / Conductor').font.size = Pt(12)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 1. RESUMEN EJECUTIVO
# ══════════════════════════════════════════════════════════════════════════════
h1('1. Resumen Ejecutivo')
body(
    'Sesión enfocada completamente en el módulo de Rutas / Muelle / Conductor. '
    'Se corrigieron bugs críticos de UX, se implementó el flujo completo de entrega '
    'parcial con detalle de referencias por producto, y se normalizó la paleta de '
    'colores al modo día en el módulo del conductor. Total: 20 commits en producción (Railway).'
)
separador()

# ══════════════════════════════════════════════════════════════════════════════
# 2. ARQUITECTURA DEL FLUJO DE RUTAS
# ══════════════════════════════════════════════════════════════════════════════
h1('2. Arquitectura del Flujo de Rutas (estado final)')

h2('2.1 Estados de RutaDespacho')
tabla(
    ['Estado', 'Quién lo activa', 'Acción del admin'],
    [
        ['PROGRAMADO',  'Admin crea ruta',         'Espera bultos disponibles'],
        ['EN_CARGUE',   'Admin "Iniciar Cargue"',   'Asigna bultos en muelle manualmente'],
        ['EN_TRANSITO', 'Admin "Salió"',             'Ve "🚛 En camino" (no puede marcar entregada)'],
        ['ENTREGADA',   'Conductor "Cerrar Ruta"',   'Abre Planilla / Liquida'],
        ['LIQUIDADA',   'Admin "Liquidar"',          'Flujo cerrado'],
    ]
)
separador()

h2('2.2 Flujo paso a paso')
pasos = [
    'Admin crea ruta desde ruta maestra (conductor + vehículo + fecha) → PROGRAMADO',
    'Admin presiona "Iniciar Cargue" → EN_CARGUE → redirige al Muelle',
    'Muelle: admin asigna bultos MANUALMENTE con botones "+ Todo el pedido" o "+ Solo esta" (sin auto-asignación)',
    'Admin escanea cada bulto físico → estado CARGADO',
    'Admin presiona "Salió" → EN_TRANSITO',
    'Conductor confirma paradas desde su app (ENTREGADO / PARCIAL / RECHAZADO)',
    'Conductor presiona "Cerrar Ruta" → ENTREGADA',
    'Admin abre Planilla → revisa resumen financiero → "Liquidar" → LIQUIDADA',
]
for i, p in enumerate(pasos, 1):
    bullet(f'{i}. {p}')
separador()

# ══════════════════════════════════════════════════════════════════════════════
# 3. MÓDULO CONDUCTOR
# ══════════════════════════════════════════════════════════════════════════════
h1('3. Módulo Conductor')

h2('3.1 Reglas de validación por estado de entrega')
tabla(
    ['Campo', 'ENTREGADO', 'PARCIAL', 'RECHAZADO'],
    [
        ['Forma de pago',        '✓ Obligatoria', '✓ Obligatoria',  '— Oculta'],
        ['Monto cobrado',        '✓ Obligatorio', '✓ > 0',          '— Oculto'],
        ['Observaciones',        '— Opcional',    '✓ Obligatoria',  '✓ OBLIGATORIA'],
        ['Bultos rechazados',    '— No aplica',   '✓ Seleccionar',  '✓ Seleccionar'],
        ['Detalle referencias',  '— No aplica',   '✓ Por ítem',     '— No aplica'],
        ['Foto evidencia',       '— Opcional',    '— Opcional',     '— Opcional'],
    ]
)
separador()

h2('3.2 Entrega Parcial — Detalle de referencias')
body(
    'Cuando el conductor selecciona PARCIAL, aparece la sección "REFERENCIAS ENTREGADAS" '
    'con un card por cada ítem del pedido:'
)
bullet('PEDIDO: cantidad total comprometida')
bullet('ENTREGADO: campo editable (0 … pedido)')
bullet('DEVUELTO: calculado automáticamente = pedido - entregado')
bullet('Validaciones: no puede confirmar sin al menos una devolución real; no puede confirmar sin observaciones')
separador()

h2('3.3 Foto con cámara directa')
body('El botón "📷 Tomar foto con la cámara" reemplaza el input de archivo nativo:')
bullet('Abre directamente la cámara trasera en Android/iOS (capture=environment)')
bullet('Muestra previsualización inmediata tras tomar la foto')
bullet('Botón "✕ Quitar foto" para volver a tomar')
bullet('En escritorio: fallback al selector de archivos del sistema operativo')
separador()

h2('3.4 Timer de refresco (30s) — comportamiento')
tabla(
    ['Estado del conductor', 'Acción del timer'],
    [
        ['Llenando formulario de parada', 'No hace nada — no toca el DOM'],
        ['Viendo lista de paradas',        'Refresca la lista silenciosamente'],
        ['Viendo lista de rutas',          'Comportamiento normal'],
    ]
)
separador()

# ══════════════════════════════════════════════════════════════════════════════
# 4. MÓDULO ADMIN — RUTAS/VIAJES
# ══════════════════════════════════════════════════════════════════════════════
h1('4. Módulo Administrador — Rutas / Viajes')

h2('4.1 Modal "Ver" (📋)')
body('Muestra el estado de cada PD en la ruta:')
bullet('Badge de color: ✓ Entregado (verde) / ⚠ Parcial (amarillo) / ✗ Rechazado (rojo) / Sin gestionar (gris)')
bullet('Bultos con indicador ✓/✗ por código de barras')
bullet('PARCIAL: grid REFERENCIA | PEDIDO | ENTREGADO | DEVUELTO')
bullet('RECHAZADO: motivo en texto + botón "📷 Ver evidencia fotográfica" → abre imagen en nueva pestaña (descargable)')
separador()

h2('4.2 Modal "Planilla de Cuadre" (💰)')
body('Resumen financiero de la ruta:')
bullet('Totales por forma de pago (Efectivo, Transferencia, Cheque, Crédito, Exento)')
bullet('Total recaudado')
bullet('Por parada: estado, monto, forma de pago, detalle PARCIAL por referencia, motivo RECHAZADO en texto')
bullet('Sin foto inline — la evidencia solo se consulta desde el modal Ver')
bullet('Botón "Liquidar" aparece cuando todas las paradas están gestionadas')
separador()

h2('4.3 Indicador de estado EN_TRANSITO')
body(
    'Mientras la ruta está EN_TRANSITO, el admin ve "🚛 En camino" '
    '(no clickeable, solo informativo). La ruta solo pasa a ENTREGADA '
    'cuando el conductor confirma desde su app. El admin conserva "⚡ Forzar cierre" '
    'como acción de emergencia.'
)
separador()

h2('4.4 Muelle — asignación manual de bultos')
body(
    'Al presionar "Iniciar Cargue", la ruta queda vacía de bultos. '
    'El admin asigna manualmente desde el Monitor de Muelle:'
)
bullet('"+ Todo el pedido" → asigna todos los bultos de un PD a la ruta activa')
bullet('"+ Solo esta" → asigna un bulto específico')
bullet('"×" → desasigna un bulto de la ruta')
body(
    'Eliminada la auto-asignación masiva que antes cargaba todos los '
    'bultos del municipio al iniciar cargue.'
)
separador()

# ══════════════════════════════════════════════════════════════════════════════
# 5. MODELO DE DATOS — CAMBIOS
# ══════════════════════════════════════════════════════════════════════════════
h1('5. Cambios en Modelo de Datos')

h2('5.1 RecaudoEntrega — nueva columna items_entregados')
body('Migración: f8e9d0c1b2a3 — aplicada en Railway ✅')
body('Columna: items_entregados JSON (nullable)')
body('Estructura del JSON para paradas PARCIAL:')

code_p = doc.add_paragraph()
code_run = code_p.add_run(
    '[\n'
    '  {\n'
    '    "codigo": "PAPELSP9218",\n'
    '    "nombre": "Resma Papel Carta",\n'
    '    "unidad": "und",\n'
    '    "cantidad_pedida": 4,\n'
    '    "cantidad_entregada": 2,\n'
    '    "cantidad_devuelta": 2\n'
    '  }\n'
    ']'
)
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)
separador()

# ══════════════════════════════════════════════════════════════════════════════
# 6. BUGS CORREGIDOS
# ══════════════════════════════════════════════════════════════════════════════
h1('6. Bugs Corregidos (20 commits en producción)')

tabla(
    ['#', 'Commit', 'Descripción'],
    [
        ['1',  'ff82475', 'Escáner español envía apóstrofe en lugar de guion → normalización en muelle_service.py'],
        ['2',  '6c1520e', 'Campos undefined en feedback de scan (nombre incorrecto del campo backend)'],
        ['3',  'dc7560b', 'window.alert con HTML crudo → modal propio en "Ver manifiesto"'],
        ['4',  '0ec1b70', 'Timer 30s sacaba al conductor de la vista de paradas'],
        ['5',  '3b237c3', 'Timer 30s pisaba formulario de confirmación de parada'],
        ['6',  '6d19550', 'Botones resultado entrega en dark mode → paleta modo día'],
        ['7',  '05b5ef5', 'Admin podía marcar ruta como entregada → reemplazado por "En camino"'],
        ['8',  '38e82f8', 'feat: entrega parcial con detalle de referencias (migración incluida)'],
        ['9',  'cc95ab8', 'Conflicto de ID de migración Alembic → corregido'],
        ['10', '3d20d36', 'Ocultar forma de pago y monto en estado RECHAZADO'],
        ['11', '5a9382f', 'feat: detalle completo de paradas RECHAZADO en planilla admin'],
        ['12', '645bcb4', 'AttributeError NoneType.upper() cuando forma_pago es null'],
        ['13', '7f7e9d9', 'feat: status de entrega por PD en modal Ver de ruta'],
        ['14', '6fa6079', 'feat: validación explícita por campo en confirmar_parada'],
        ['15', '1d4575a', 'Modal Ver manifiesto → paleta modo día completa'],
        ['16', 'df1eee1', 'Eliminar auto-asignación de bultos al iniciar cargue'],
        ['17', '7d8c84f', 'Tarjetas de paradas del conductor → paleta modo día'],
        ['18', '9a73065', 'Observaciones obligatorias en RECHAZADO + evidencia como botón'],
        ['19', '2813cc0', 'feat: botón cámara directa en confirmación de parada'],
        ['20', 'f4a2f27', 'Eliminar evidencia fotográfica de Planilla de Cuadre'],
    ]
)
separador()

# ══════════════════════════════════════════════════════════════════════════════
# 7. PENDIENTES PRÓXIMA SESIÓN
# ══════════════════════════════════════════════════════════════════════════════
h1('7. Pendientes para Próxima Sesión')

h2('7.1 🔴 CRÍTICO — Devoluciones de ruta → módulo recepción')
body('Análisis completo hecho. Pendiente implementar. Requiere autorización para migración.')
body('Plan (Opción A — reutilizar TareaDevolucion):')
bullet('Agregar columnas origen_tipo y origen_ruta_id a tabla tareas_devolucion')
bullet('PARCIAL: al confirmar parada, crear TareaDevolucion por cada referencia con cantidad_devuelta > 0')
bullet('RECHAZADO: al confirmar parada, crear TareaDevolucion por cada bulto rechazado')
bullet('Recepcionista gestiona desde tab Devoluciones (flujo ya conocido: ubicar / averiado / descartar)')
separador()

h2('7.2 🟡 BLOQUEADO — Picking parcial (consulta 7811)')
body('El usuario del Módulo de Conectividad SQL Server no tiene GRANT UPDATE sobre dbo.t405_cm_compromisos.')
body('Ticket enviado a Siesa Soporte. Código WMS listo en DIAGNOSTICO_PICKING_PARCIAL.md.')
bullet('Acción: preguntar al inicio de la sesión si Siesa dio el permiso')
bullet('Si sí → activar en 5 minutos')
separador()

h2('7.3 🟡 PENDIENTE — Recepción OC (conector 142948)')
body('Fix Tipo de Proveedor 0001 para NIT 1001196008 pendiente de confirmar por consultor Siesa.')
bullet('NO re-ejecutar prueba hasta confirmación del consultor')
separador()

h2('7.4 🔵 CONFIGURACIÓN — Backorder Siesa (sin código)')
body('Duglas debe configurar Backorder=1 para todos los clientes críticos:')
bullet('Maestros → Clientes → NIT → Sucursales → Ventas → Backorder = "1 — Despachar solo disponible y cancelar resto"')
bullet('Sin esto, conector 142945 ignora la cantidad real del WMS y usa la cantidad comprometida del pedido')
separador()

h2('7.5 🔵 MEJORA — Bug zona vs tipo_zona en put-away')
body('_buscar_ubicacion_optima filtra por campo zona pero el modelo Ubicacion usa tipo_zona.')
body('Resultado: put-away queda sin ubicación física. Pendiente investigar y corregir.')
separador()

# ══════════════════════════════════════════════════════════════════════════════
# 8. REFERENCIA DE ARCHIVOS CLAVE
# ══════════════════════════════════════════════════════════════════════════════
h1('8. Referencia de Archivos Clave')

tabla(
    ['Archivo', 'Responsabilidad'],
    [
        ['app/services/muelle_service.py',    'Scan, carga y asignación de bultos en muelle'],
        ['app/services/ruta_service.py',       'Rutas, paradas, confirmar_parada, planilla, liquidación'],
        ['app/models/recaudo_entrega.py',      'Modelo de confirmación de entrega (incluye items_entregados)'],
        ['app/routes/muelle.py',               'Endpoints HTTP del muelle'],
        ['app/routes/rutas.py',                'Endpoints HTTP de rutas'],
        ['app/static/pwa/app.js',              'PWA completa — módulo conductor, admin muelle/rutas, planilla'],
        ['app/services/devolucion_service.py', 'Devoluciones (base para integrar devoluciones de ruta)'],
        ['DIAGNOSTICO_PICKING_PARCIAL.md',     'Diagnóstico bug 142945 + código listo para activar'],
        ['SESION_2026-06-01.md',               'Este documento en formato Markdown'],
    ]
)
separador()

# ══════════════════════════════════════════════════════════════════════════════
# 9. VARIABLES RAILWAY
# ══════════════════════════════════════════════════════════════════════════════
h1('9. Variables Railway Confirmadas')

tabla(
    ['Variable', 'Valor', 'Estado'],
    [
        ['SIESA_TIPO_DOCTO_REMISION',         'RM',   '✅ Activa'],
        ['SIESA_ID_MOTIVO_VENTAS',            '01',   '✅ Activa'],
        ['SIESA_BODEGA_TRANSITO',             'TRA1', '✅ Activa'],
        ['SIESA_TIPO_DOCTO_TRANSITO_SALIDA',  'STS',  '✅ Activa'],
        ['SIESA_TIPO_DOCTO_TRANSITO_ENTRADA', 'TTE',  '✅ Activa'],
        ['SIESA_UNIDAD_NEGOCIO',              '001',  '✅ Activa'],
        ['SIESA_PUNTO_ENVIO_DEFAULT',         '000',  '✅ Activa'],
    ]
)

# Pie de página
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run(f'Generado: {datetime.now().strftime("%Y-%m-%d %H:%M")}  ·  WMS-PAME  ·  CTO Claude Sonnet 4.6').font.size = Pt(9)

doc.save('SESION_2026-06-01.docx')
print('OK - SESION_2026-06-01.docx generado')
