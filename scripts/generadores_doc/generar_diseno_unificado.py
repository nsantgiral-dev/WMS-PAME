"""Genera el documento de diseño definitivo de Picking/Packing Unificado."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Estilos globales ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

def set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def body(doc, text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    # fondo gris claro con shading en el párrafo
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)
    return p

def add_table(doc, headers, rows, col_widths=None, header_fill='1F497D'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr_cells[i], header_fill)
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].paragraphs[0].add_run(val).font.size = Pt(9.5)
    if col_widths:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    set_col_width(cell, col_widths[j])
    return table

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('WMS-PAME')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Diseño Definitivo')
run2.bold = True
run2.font.size = Pt(20)
run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Unificación de Picking y Packing')
run3.font.size = Pt(16)
run3.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run('Pedidos (PD) y Requisiciones / Traslados (ST)')
run4.font.size = Pt(13)
run4.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()
doc.add_paragraph()
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run5 = p5.add_run('Versión 1.0  |  Junio 2026')
run5.font.size = Pt(10)
run5.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
run6 = p6.add_run('Principios SOLID  |  Strategy Pattern  |  Multi-Bodega')
run6.font.size = Pt(10)
run6.italic = True
run6.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. CONTEXTO Y OBJETIVO
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '1. Contexto y Objetivo')

body(doc, 'El WMS-PAME gestiona hoy dos tipos de solicitudes de despacho que comparten la '
          'misma lógica operativa de preparación pero divergen en el cierre contable en SIESA:')

add_table(doc,
    ['Tipo', 'Código', 'Origen', 'Destino', 'Cierre en SIESA'],
    [
        ['Pedido de Venta',         'PD',   'NB1 (CO003)',    'Cliente externo',     'Factura (FE) + Remisión (RM)'],
        ['Requisición / Traslado',  'ST',   'Cualquier bodega', 'Otra bodega/tienda', 'RIT → Compromisos → STS → ETS'],
    ],
    col_widths=[4, 2, 3.5, 3.5, 5]
)

doc.add_paragraph()
body(doc, 'El objetivo es unificar el módulo de Picking y Packing para que el operario '
          'trabaje con una sola pantalla para ambos tipos, mientras el sistema resuelve '
          'internamente qué documentos genera en SIESA al momento del cierre. '
          'Adicionalmente, el flujo debe soportar N tiendas sin cambios de código.')

# ══════════════════════════════════════════════════════════════════════════════
# 2. APIS SIESA
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '2. APIs SIESA para Traslados')

add_table(doc,
    ['Conector', 'Tipo', 'Nombre API', 'Cuándo se dispara'],
    [
        ['174646', 'POST escritura', 'RequisicionesParaTransferir (RIT)',   'Al confirmar picking completado'],
        ['174720', 'POST escritura', 'CompromisosDesdeRequisicion',         'Paso 1 al cerrar caja ST'],
        ['174930', 'POST escritura', 'TransferenciasDesdeRequisicion (STS)','Paso 2 al cerrar caja ST'],
        ['173079', 'POST escritura', 'TransferenciaEnTransitoEntrada (ETS)','Tienda destino confirma recepción'],
        ['176',    'GET consulta',   'Transferencia_Salida_Transito',       'Verificar STS antes de recibir'],
        ['177',    'GET consulta',   'Transferencia_Transito_Entrada',      'Verificar ETS creada'],
        ['179',    'GET consulta',   'Transferencias general',              'Consultas de estado generales'],
    ],
    col_widths=[2, 2.8, 6.5, 5.5]
)

doc.add_paragraph()
h3(doc, 'Notas críticas de las APIs')
bullet(doc, '174930 (STS): solo requiere el consecutivo de la RIT (f440_consec_docto_req_int). '
            'SIESA resuelve los ítems internamente. No se reenvían ítems.')
bullet(doc, '174720 (Compromisos): requiere los ítems con cantidad_real del doble conteo '
            '(f441_cant_por_remisionar_base) y bodega origen + destino por ítem. '
            'Por eso se dispara al cerrar caja, no antes.')
bullet(doc, 'Los GET 176/177/179 son de consulta. Se usan para verificar estado '
            'antes de permitir la recepción en tienda destino.')

# ══════════════════════════════════════════════════════════════════════════════
# 3. MODELO DE DATOS
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '3. Modelo de Datos — Solo Extensiones')

body(doc, 'Todos los cambios son adiciones. Ningún campo existente se modifica ni elimina. '
          'Los registros actuales reciben backfill automático en la migración.')

h3(doc, '3.1  TareaPicking  (tabla: tareas_picking)')
code_block(doc, '+ bodega_origen_siesa  VARCHAR(20)   nullable\n'
                '  Valores: \'NB1\', \'NC1\', \'NS1\', \'NT1\'... (cualquier bodega SIESA)\n'
                '  Migration backfill: UPDATE tareas_picking SET bodega_origen_siesa = \'NB1\'')

h3(doc, '3.2  TareaPacking  (tabla: tareas_packing)')
code_block(doc, '+ tipo_documento       VARCHAR(20)   NOT NULL  DEFAULT \'PEDIDO\'\n'
                '  Valores: \'PEDIDO\' | \'TRASLADO\'\n\n'
                '+ referencia_doc       VARCHAR(50)   nullable\n'
                '  Ejemplo: \'PD1307\' o \'ST-20260603-001\'\n\n'
                '+ solicitud_id         INTEGER       FK → solicitudes_traslado  NULL\n'
                '  Solo se popula cuando tipo_documento = \'TRASLADO\'\n\n'
                '+ tienda_destino       VARCHAR(100)  nullable\n'
                '  Nombre display en UI: \'Tienda Centro\', \'Tienda Sur\'...\n\n'
                '+ bodega_origen_siesa  VARCHAR(20)   nullable\n'
                '  Migration backfill: UPDATE tareas_packing SET bodega_origen_siesa = \'NB1\'')

h3(doc, '3.3  SiesaJob  (tabla: siesa_jobs)')
code_block(doc, '+ compromisos_ok   BOOLEAN  DEFAULT FALSE\n'
                '  TRUE = paso 174720 ya ejecutado; el retry lo omite\n\n'
                '+ consec_sts       INTEGER  NULL\n'
                '  Consecutivo STS devuelto por 174930; se guarda para consultas\n\n'
                '+ sts_ok           BOOLEAN  DEFAULT FALSE\n'
                '  TRUE = paso 174930 ya ejecutado; el retry lo omite\n\n'
                'Tipo nuevo: \'DESPACHO_TRASLADO\'\n'
                '  Payload: {solicitud_id, consec_rit, bodega_origen, bodega_destino,\n'
                '            items: [{id, cant_real, bodega_sal, bodega_ent}]}')

h3(doc, '3.4  Usuario  (sin cambios)')
body(doc, 'Ya posee almacen_id (para NB1) y bodega_siesa_id (para tiendas). '
          'El scoping se resuelve con estos dos campos existentes.')

h3(doc, '3.5  SolicitudTraslado  (sin cambios)')
body(doc, 'Ya posee bodega_origen_siesa y bodega_destino_siesa. '
          'Soporta cualquier combinación bodega → bodega sin cambios.')

doc.add_paragraph()
h3(doc, 'Migraciones requeridas (requieren autorización explícita)')
add_table(doc,
    ['#', 'Migración', 'Detalle'],
    [
        ['005', 'ADD bodega_origen_siesa',                'tareas_picking + backfill NB1'],
        ['006', 'ADD tipo_documento, referencia_doc,\nsolicitud_id, tienda_destino,\nbodega_origen_siesa', 'tareas_packing + backfill'],
        ['007', 'ADD compromisos_ok, consec_sts, sts_ok\n+ tipo DESPACHO_TRASLADO', 'siesa_jobs'],
    ],
    col_widths=[1.2, 5.5, 5.5]
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. ARQUITECTURA DE SERVICIOS — SOLID
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '4. Arquitectura de Servicios — Principios SOLID')

h3(doc, '4.1  Estructura de archivos')
code_block(doc,
    'app/services/\n'
    '  closing/\n'
    '    base.py              ← IPackingCloser (ABC)\n'
    '                              def ejecutar_cierre(tarea: TareaPacking) → CierreResult\n'
    '    pedido_closer.py     ← PedidoPackingCloser(IPackingCloser)\n'
    '                              lógica 238925: Factura + Remisión\n'
    '    traslado_closer.py   ← TrasladoPackingCloser(IPackingCloser)\n'
    '                              lógica 174720 + 174930 con idempotencia\n'
    '    factory.py           ← PackingCloserFactory\n'
    '                              .get(\'PEDIDO\')   → PedidoPackingCloser\n'
    '                              .get(\'TRASLADO\') → TrasladoPackingCloser\n\n'
    '  scoping/\n'
    '    task_scope.py        ← _scope_picking(user, query)\n'
    '                           _scope_packing(user, query)\n'
    '                              un solo lugar, una sola responsabilidad\n\n'
    '  picking_service.py     ← agrega bodega_origen_siesa al crear tareas\n'
    '  packing_service.py     ← cerrar_packing() ~30 líneas, delega al factory\n'
    '  traslado_service.py    ← confirmar_picking crea TareaPacking\n'
    '                           se elimina confirmar_packing_traslado()')

h3(doc, '4.2  SOLID aplicado')
add_table(doc,
    ['Principio', 'Aplicación'],
    [
        ['S — Single Responsibility',
         'Cada closer tiene una sola razón para cambiar: su tipo de documento en SIESA. '
         'task_scope.py centraliza toda la lógica de filtrado.'],
        ['O — Open/Closed',
         'Agregar tipo DEVOLUCION = nuevo archivo DevolucionPackingCloser + una línea en factory. '
         'PackingService.cerrar_packing() no se toca nunca más.'],
        ['L — Liskov Substitution',
         'PedidoPackingCloser y TrasladoPackingCloser son intercambiables '
         'donde se use IPackingCloser. Mismo contrato, distinto comportamiento.'],
        ['I — Interface Segregation',
         'IPackingCloser expone solo ejecutar_cierre(). '
         'Tienda users acceden solo a endpoints de TRASLADO, nunca a PD.'],
        ['D — Dependency Inversion',
         'PackingService.cerrar_packing() depende de IPackingCloser (abstracción). '
         'La factory resuelve las implementaciones concretas, no el servicio.'],
    ],
    col_widths=[4, 12.5]
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. SCOPING MULTI-BODEGA
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '5. Scoping Multi-Bodega — N Tiendas')

body(doc, 'El diseño es agnóstico al número de tiendas. No existe lista de bodegas '
          'hardcodeada. Agregar una tienda nueva equivale a crear un usuario con el '
          'bodega_siesa_id correspondiente. Cero cambios de código.')

h3(doc, '5.1  Lógica de scoping centralizada')
body(doc, 'Un único punto de decisión. NB1 ve PD y ST. Tiendas ven solo ST de su bodega.')
code_block(doc,
    '_scope_picking(user, query):\n'
    '  si user.rol in GESTION:\n'
    '    return query                                    # admin ve todo\n'
    '  si user.rol == TIENDA:\n'
    '    return query.filter(\n'
    '      bodega_origen_siesa = user.bodega_siesa_id,  # solo SU bodega\n'
    '      tipo_documento = \'TRASLADO\'                   # NUNCA ve PD — solo tiendas\n'
    '    )\n'
    '  else:  # operario / empacador NB1\n'
    '    # NB1 ve PD y ST sin filtro de tipo_documento\n'
    '    return query.filter(almacen_id = user.almacen_id)\n\n'
    '_scope_packing(user, query):  # misma lógica exacta\n'
    '  # La restricción tipo_documento=TRASLADO para tiendas\n'
    '  # garantiza que PD nunca aparezca fuera de NB1')

h3(doc, '5.2  Tabla de visibilidad por rol')
body(doc, 'Regla fundamental: NB1 (CO003) es la ÚNICA bodega que puede ver y gestionar '
          'Pedidos (PD). Todas las demás tiendas ven exclusivamente Traslados (ST). '
          'Esta restricción aplica tanto en picking como en packing.')
add_table(doc,
    ['Usuario', 'Bodega', 'Picking — ve', 'Packing — ve', 'Traslados', 'PD', 'Conteo'],
    [
        ['Admin / Supervisor', 'NB1',   'PD + ST de NB1',      'PD + ST de NB1',      'Todos',               'Sí', 'Sí'],
        ['Operario',           'NB1',   'PD + ST de NB1',      'No',                  'No',                  'Sí', 'Sí'],
        ['Empacador',          'NB1',   'No',                  'PD + ST de NB1',      'No',                  'Sí', 'No'],
        ['Tienda NC1',         'NC1',   'Solo ST origen NC1',  'Solo ST origen NC1',  'Origen o destino NC1','No', 'No'],
        ['Tienda NS1',         'NS1',   'Solo ST origen NS1',  'Solo ST origen NS1',  'Origen o destino NS1','No', 'No'],
        ['Tienda NXX',         'NXX',   'Solo ST origen NXX',  'Solo ST origen NXX',  'Origen o destino NXX','No', 'No'],
    ],
    col_widths=[3.2, 1.8, 3.3, 3.3, 3.8, 1, 1.3]
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. PRIORIDAD Y BLOQUEO EN PICKING
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '6. Prioridad y Bloqueo en Picking')

h3(doc, '6.1  Cola de asignación')
body(doc, 'La cola de prioridad aplica diferente según la bodega del operario.')
add_table(doc,
    ['Prioridad', 'Tipo de tarea', 'Aplica a', 'Condición'],
    [
        ['1°', 'Pedido (PD)',      'NB1 únicamente',    'TareaPicking PENDIENTE, tipo=PEDIDO, almacen=NB1'],
        ['2°', 'Requisición (ST)', 'NB1 y todas las tiendas', 'TareaPicking PENDIENTE, tipo=TRASLADO, bodega del usuario — NB1 solo si no hay PD'],
        ['3°', 'Conteo Cíclico',   'NB1 únicamente',    'TareaConteo PENDIENTE, almacen=NB1 — solo si no hay PD ni ST'],
    ],
    col_widths=[1.8, 3.2, 4.5, 7]
)
doc.add_paragraph()
body(doc, 'Las tiendas (NC1, NS1, NXX...) solo tienen la prioridad 2° (ST). '
          'Nunca ven PD ni Conteo Cíclico. Su cola siempre es de Traslados.')

h3(doc, '6.2  Regla de bloqueo — dos capas')
body(doc, 'Una vez que un operario tiene una tarea EN_PROCESO no puede tomar otra, '
          'independientemente de su prioridad. El bloqueo aplica en backend y frontend.')

code_block(doc,
    '# Backend — POST /api/picking/{id}/iniciar\n'
    'tarea_activa = TareaPicking.query.filter_by(\n'
    '    operario_id = user.id,\n'
    '    estado = \'EN_PROCESO\'\n'
    ').first()\n\n'
    'if tarea_activa:\n'
    '    → HTTP 409 + id de la tarea activa\n'
    '    # El frontend abre directamente la tarea activa\n\n'
    '# Frontend — cuando operario tiene tarea EN_PROCESO:\n'
    '    → Muestra solo esa tarea, sin opción de cambiar')

h3(doc, '6.3  Eventualidades resueltas')
add_table(doc,
    ['Situación', 'Resolución'],
    [
        ['Operario con ST activo, llega PD urgente',
         'PD queda en posición 1 de la cola. El operario no lo ve hasta completar el ST. '
         'Al liberar, el sistema le muestra el PD automáticamente.'],
        ['Operario con PD activo, llega ST',
         'ST queda en posición 2 de la cola. El operario termina el PD primero. '
         'Sin intervención humana.'],
        ['No hay PD ni ST pendientes',
         'El operario ve tareas de Conteo Cíclico (solo NB1). '
         'Si tampoco hay conteo, la cola aparece vacía.'],
        ['Dos operarios libres, un solo PD',
         'El primero en presionar "iniciar" adquiere el lock. '
         'El segundo recibe HTTP 409 y obtiene el siguiente ítem disponible.'],
    ],
    col_widths=[5.5, 11]
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. FLUJO PEDIDO (PD) — COMPLETO
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '7. Flujo Pedido (PD) — Completo')

body(doc, 'Exclusivo de NB1 (CO003). Genera Factura Electrónica + Remisión en SIESA.')

h3(doc, '7.1  Diagrama de flujo')
code_block(doc,
    'SIESA sync (5 min) → PedidoSiesa en BD\n'
    '│\n'
    'ADMIN NB1\n'
    '│  POST /api/siesa/iniciar-despacho\n'
    '│  └→ PickingService.crear_tareas()\n'
    '│       TareaPicking[]: tipo=PEDIDO, almacen_id=NB1,\n'
    '│                       bodega_origen_siesa=\'NB1\'\n'
    '│       Estado: PENDIENTE\n'
    '│\n'
    'OPERARIO NB1  (cola muestra PD — prioridad 1)\n'
    '│  PUT  /api/picking/{id}/iniciar         → EN_PROCESO (lock activo)\n'
    '│  POST /api/picking/{id}/confirmar-cantidad\n'
    '│       cantidad_recogida = X  → COMPLETADO\n'
    '│       Descuenta stock UbicacionProducto\n'
    '│       Lock liberado\n'
    '│\n'
    'ADMIN NB1\n'
    '│  POST /api/packing/crear-desde-picking\n'
    '│  └→ PackingService.crear_desde_picking()\n'
    '│       TareaPacking: tipo=PEDIDO, referencia=\'PD1307\',\n'
    '│                     almacen_id=NB1, bodega_origen_siesa=\'NB1\'\n'
    '│       ItemPacking[]: cantidad_esperada = cantidad_recogida\n'
    '│\n'
    'EMPACADOR NB1  (cola muestra etiqueta [PEDIDO])\n'
    '│  PUT  /api/packing/{id}/iniciar         → EN_PROCESO\n'
    '│  POST /api/packing/{id}/escanear-item × N   ← doble conteo real\n'
    '│  PUT  /api/packing/{id}/confirmar       → VERIFICADO\n'
    '│  POST /api/packing/{id}/cerrar  {bultos: [{tipo, cantidad}]}\n'
    '│  └→ PackingService.cerrar_packing()\n'
    '│       factory.get(\'PEDIDO\') → PedidoPackingCloser\n'
    '│       Crea Bulto[]: PD1307-01, PD1307-02...\n'
    '│       Enqueue SiesaJob(tipo=DESPACHO_F470)\n'
    '│       Commit atómico bultos + job\n'
    '│       TareaPacking → DESPACHADO\n'
    '│\n'
    'DLQ RAILWAY (async — retry automático cada 5 min)\n'
    '   SiesaJob DESPACHO_F470\n'
    '   └→ Connekta 238925 FacturaPedido\n'
    '        → Factura Electrónica (FE) + Remisión (RM)\n'
    '        siesa_triggered = True\n'
    '        RM imprimible: GET /api/packing/{id}/remision')

h3(doc, '7.2  Pedido parcial')
body(doc, 'Si el picking quedó incompleto (ítems en estado BLOQUEADO), el packing se crea '
          'con lo efectivamente recogido. PedidoPackingCloser detecta que '
          'cantidad_pendiente > 0 y activa el flujo de despacho parcial existente '
          '(244328 → 142945 → 142943). Ese flujo no recibe ningún cambio.')

# ══════════════════════════════════════════════════════════════════════════════
# 8. FLUJO REQUISICIÓN / TRASLADO (ST) — COMPLETO
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '8. Flujo Requisición / Traslado (ST) — Completo')

body(doc, 'Aplica para cualquier combinación de bodegas: NB1→Tienda, Tienda→Tienda, '
          'Tienda→NB1. El ciclo SIESA es: RIT (174646) → Compromisos (174720) → '
          'STS (174930) → ETS (173079).')

h3(doc, '8.1  Diagrama de flujo')
code_block(doc,
    'USUARIO DE BODEGA ORIGEN  (NB1, NC1, NS1... cualquier bodega)\n'
    '│  POST /api/traslados/\n'
    '│       bodega_origen_siesa = \'NC1\'\n'
    '│       bodega_destino_siesa = \'NS1\'\n'
    '│       items: [{producto_id, cantidad_solicitada}]\n'
    '│       Estado: BORRADOR\n'
    '│  POST /api/traslados/{id}/enviar  → ENVIADA\n'
    '│\n'
    'ADMIN  (aprueba)\n'
    '│  POST /api/traslados/{id}/aprobar\n'
    '│  └→ TrasladoService.aprobar_solicitud()\n'
    '│       TareaPicking[]: tipo=TRASLADO,\n'
    '│                       bodega_origen_siesa=\'NC1\',\n'
    '│                       almacen_id=NULL\n'
    '│       Estado traslado: EN_PICKING\n'
    '│\n'
    'OPERARIO DE NC1  (cola muestra ST si no hay PD activo — prioridad 2)\n'
    '│  Mismo flujo de conteo ciego que PD\n'
    '│  Al completar último ítem:\n'
    '│  └→ TrasladoService.confirmar_picking_traslado()\n'
    '│       POST Connekta 174646 (RIT)\n'
    '│          f440_id_bodega_salida = \'NC1\'\n'
    '│          f440_id_bodega_entrada = \'NS1\'\n'
    '│          f441_id_item + f441_cant_base por ítem\n'
    '│       → siesa_requisicion_consec guardado\n'
    '│       Crea TareaPacking: tipo=TRASLADO,\n'
    '│                          referencia=\'ST-20260603-001\',\n'
    '│                          solicitud_id=X,\n'
    '│                          tienda_destino=\'Tienda Sur\',\n'
    '│                          bodega_origen_siesa=\'NC1\'\n'
    '│       ItemPacking[]: cantidad_esperada = cantidad_recogida\n'
    '│       Estado traslado: EN_PACKING\n'
    '│\n'
    'EMPACADOR DE NC1  (cola: etiqueta [TRASLADO] + "→ Tienda Sur")\n'
    '│  PUT  /api/packing/{id}/iniciar         → EN_PROCESO\n'
    '│  POST /api/packing/{id}/escanear-item × N   ← MISMO doble conteo que PD\n'
    '│  PUT  /api/packing/{id}/confirmar       → VERIFICADO\n'
    '│  POST /api/packing/{id}/cerrar  {bultos: [{tipo, cantidad}]}\n'
    '│  └→ PackingService.cerrar_packing()\n'
    '│       factory.get(\'TRASLADO\') → TrasladoPackingCloser\n'
    '│       Crea Bulto[]\n'
    '│       Enqueue SiesaJob(tipo=DESPACHO_TRASLADO,\n'
    '│                        consec_rit = siesa_requisicion_consec,\n'
    '│                        bodega_origen = \'NC1\',\n'
    '│                        bodega_destino = \'NS1\',\n'
    '│                        items = [{id, cant_real, bodega_sal, bodega_ent}])\n'
    '│       Commit atómico bultos + job\n'
    '│       TareaPacking → DESPACHADO\n'
    '│\n'
    'DLQ RAILWAY  (async — idempotencia de 2 pasos)\n'
    '│\n'
    '│  SiesaJob DESPACHO_TRASLADO\n'
    '│  │\n'
    '│  ├─ PASO 1  (si compromisos_ok = False):\n'
    '│  │    POST Connekta 174720 (Compromisos desde RIT)\n'
    '│  │       f440_consec_docto = consec_rit\n'
    '│  │       por ítem: f441_cant_base = cantidad_real empacada\n'
    '│  │                 f441_id_bodega = bodega_salida\n'
    '│  │                 f441_id_bodega_ent = bodega_destino\n'
    '│  │                 f441_cant_por_remisionar_base = cantidad_real\n'
    '│  │    → OK: compromisos_ok = True  (commit inmediato)\n'
    '│  │\n'
    '│  └─ PASO 2  (si sts_ok = False):\n'
    '│       POST Connekta 174930 (Transfer desde RIT → STS)\n'
    '│          f440_id_co_req_int = CO origen\n'
    '│          f440_id_tipo_docto_req_int = tipo RIT\n'
    '│          f440_consec_docto_req_int = consec_rit\n'
    '│       → consec_sts guardado, sts_ok = True  (commit)\n'
    '│\n'
    '│  Si falla cualquier paso → job FALLIDO → retry en 5 min\n'
    '│  El retry retoma desde el último paso fallido (idempotencia garantizada)\n'
    '│  SolicitudTraslado → EN_TRANSITO\n'
    '│\n'
    'TIENDA DESTINO (NS1 — solo ve sus traslados donde es destino)\n'
    '   WMS consulta GET 176 para verificar que STS existe\n'
    '   Operario tienda escanea ítems físicamente recibidos\n'
    '   POST /api/traslados/{id}/confirmar-recepcion\n'
    '   └→ TrasladoService.confirmar_recepcion()\n'
    '        Acepta cantidad real (completa o parcial sin bloqueo)\n'
    '        POST Connekta 173079 (ETS — Entrada en Tránsito)\n'
    '           ítems con cantidad_recibida real\n'
    '        siesa_entrada_consec guardado\n'
    '        SolicitudTraslado → ENTREGADA\n'
    '        RIT en SIESA → Estado 4 (Cumplido)')

h3(doc, '8.2  Recepción parcial')
body(doc, 'Si llegaron 9 de 10 unidades: el sistema acepta los 9, dispara ETS con '
          'cantidad_recibida=9 y avanza a ENTREGADA. La discrepancia queda registrada '
          'en el ítem (solicitado=10, recibido=9) y es visible en el panel de admin. '
          'El ajuste de inventario queda como tarea administrativa separada; '
          'no bloquea el flujo.')

h3(doc, '8.3  Máquina de estados SolicitudTraslado')
code_block(doc,
    'BORRADOR\n'
    '  └→ ENVIADA\n'
    '       ├→ RECHAZADA  (admin rechaza)\n'
    '       └→ EN_PICKING (admin aprueba + TareaPicking creada)\n'
    '            └→ EN_PACKING (picking confirmado + RIT 174646 + TareaPacking creada)\n'
    '                 └→ DESPACHADO (cerrar caja + SiesaJob DESPACHO_TRASLADO)\n'
    '                      └→ EN_TRANSITO (STS 174930 procesada por DLQ)\n'
    '                           └→ ENTREGADA (tienda destino confirma + ETS 173079)\n'
    '  └→ CANCELADA (en cualquier momento antes de EN_TRANSITO)')

# ══════════════════════════════════════════════════════════════════════════════
# 9. PANTALLA PACKING UNIFICADA
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '9. Pantalla Packing Unificada')

body(doc, 'REGLA FUNDAMENTAL: NB1 (CO003) es la ÚNICA bodega que puede ver y gestionar '
          'Pedidos (PD) en la pantalla de packing. Todas las demás tiendas ven '
          'exclusivamente Traslados (ST).')

h3(doc, '9.1  Vista NB1 — Empacador bodega principal (PD + ST)')
code_block(doc,
    '┌──────────────────────────────────────────────┐\n'
    '│  TAREAS DE EMPAQUE  [NB1 — Bodega Principal] │\n'
    '│                                              │\n'
    '│  ┌──────────────────────────────────────┐   │\n'
    '│  │ PD1307              [PEDIDO]  ●       │   │  ← etiqueta azul\n'
    '│  │ 1 producto · 0/1 verificados          │   │\n'
    '│  │ Pendiente                             │   │\n'
    '│  └──────────────────────────────────────┘   │\n'
    '│                                              │\n'
    '│  ┌──────────────────────────────────────┐   │\n'
    '│  │ ST-20260603-001   [TRASLADO]  ◆       │   │  ← etiqueta naranja\n'
    '│  │ 3 productos · 0/3 verificados         │   │\n'
    '│  │ Pendiente  →  Tienda Sur              │   │  ← muestra destino\n'
    '│  └──────────────────────────────────────┘   │\n'
    '│                                              │\n'
    '│  ┌──────────────────────────────────────┐   │\n'
    '│  │ PD1306              [PEDIDO]  ●       │   │\n'
    '│  │ 1 producto · 0/1 verificados          │   │\n'
    '│  │ Pendiente                             │   │\n'
    '│  └──────────────────────────────────────┘   │\n'
    '└──────────────────────────────────────────────┘\n'
    '  NB1 ve PD y ST mezclados. PD siempre aparece primero.')

h3(doc, '9.2  Vista Tienda — Empacador cualquier tienda (solo ST)')
code_block(doc,
    '┌──────────────────────────────────────────────┐\n'
    '│  TAREAS DE EMPAQUE  [Tienda Centro — NC1]    │\n'
    '│                                              │\n'
    '│  ┌──────────────────────────────────────┐   │\n'
    '│  │ ST-20260603-001   [TRASLADO]  ◆       │   │  ← etiqueta naranja\n'
    '│  │ 3 productos · 0/3 verificados         │   │\n'
    '│  │ Pendiente  →  Tienda Sur              │   │\n'
    '│  └──────────────────────────────────────┘   │\n'
    '│                                              │\n'
    '│  ┌──────────────────────────────────────┐   │\n'
    '│  │ ST-20260604-002   [TRASLADO]  ◆       │   │\n'
    '│  │ 5 productos · 0/5 verificados         │   │\n'
    '│  │ Pendiente  →  Tienda Norte            │   │\n'
    '│  └──────────────────────────────────────┘   │\n'
    '│                                              │\n'
    '│  (los Pedidos PD nunca aparecen aquí)        │\n'
    '└──────────────────────────────────────────────┘\n'
    '  Tiendas ven ÚNICAMENTE traslados donde son bodega origen.')

h3(doc, '9.3  Reglas de presentación')
add_table(doc,
    ['Regla', 'NB1 (CO003)', 'Tiendas (NC1, NS1, NXX...)'],
    [
        ['Qué ve en packing',   'PD + ST',                          'Solo ST'],
        ['Orden en lista',      'PD primero, luego ST (fecha ASC)',  'ST por fecha ASC'],
        ['Etiqueta PEDIDO',     'Azul — muestra número PD',         'No aparece nunca'],
        ['Etiqueta TRASLADO',   'Naranja — código ST + destino',     'Naranja — código ST + destino'],
        ['Botón Cerrar Caja',   'Idéntico para PD y ST',            'Idéntico para ST'],
        ['Pantalla de conteo',  'Idéntica para PD y ST',            'Idéntica para ST'],
    ],
    col_widths=[3.8, 5.5, 7.2]
)

# ══════════════════════════════════════════════════════════════════════════════
# 10. RESUMEN — DOS CIERRES UN SOLO BOTÓN
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '10. Resumen — Dos Cierres, Un Solo Botón')

code_block(doc,
    '            CERRAR CAJA  (mismo botón — misma pantalla)\n'
    '                              ↓\n'
    '         PackingService.cerrar_packing(tarea_id, bultos)\n'
    '                              ↓\n'
    '         PackingCloserFactory.get(tarea.tipo_documento)\n'
    '               ↙                           ↘\n'
    '          PEDIDO                        TRASLADO\n'
    '     PedidoPackingCloser           TrasladoPackingCloser\n'
    '            ↓                               ↓\n'
    '       SiesaJob                       SiesaJob\n'
    '     DESPACHO_F470                DESPACHO_TRASLADO\n'
    '            ↓                      ↓ paso 1      ↓ paso 2\n'
    '          238925                 174720          174930\n'
    '     Factura (FE)            Compromisos          STS\n'
    '     Remisión (RM)                                 ↓\n'
    '            ↓                              EN_TRANSITO\n'
    '    Cliente recibe                               ↓\n'
    '                                    Tienda destino confirma\n'
    '                                         ↓\n'
    '                                    173079 (ETS)\n'
    '                                         ↓\n'
    '                                    ENTREGADA\n'
    '                                RIT → Estado 4 Cumplido')

# ══════════════════════════════════════════════════════════════════════════════
# 11. GARANTÍA DE NO REGRESIÓN
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '11. Garantía de No Regresión')

body(doc, 'Los siguientes componentes no reciben ningún cambio. La unificación es '
          'completamente aditiva respecto al código existente.')

add_table(doc,
    ['Componente', 'Archivo', 'Estado'],
    [
        ['PickingService.calcular_fefo()',      'picking_service.py',          'Sin cambios'],
        ['PickingService.confirmar_picking()',  'picking_service.py',          'Sin cambios'],
        ['PackingService.escanear_item()',      'packing_service.py',          'Sin cambios'],
        ['PackingService.confirmar_packing()', 'packing_service.py',          'Sin cambios'],
        ['PedidoPackingCloser',                'closing/pedido_closer.py',     'Extracción 1:1 de cerrar_packing() actual'],
        ['TrasladoService.confirmar_recepcion()', 'traslado_service.py',       'Sin cambios'],
        ['Flujo despacho parcial PD',          'despacho_parcial_service.py', 'Sin cambios'],
        ['Conteo Cíclico',                     'conteo_service.py',           'Sin cambios — solo NB1'],
    ],
    col_widths=[5.5, 5.5, 5.5]
)

# ══════════════════════════════════════════════════════════════════════════════
# PIE DE PÁGINA
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_final.paragraph_format.space_before = Pt(80)
run_f = p_final.add_run('WMS-PAME  |  Diseño Definitivo Picking + Packing Unificado  |  Junio 2026')
run_f.font.size = Pt(9)
run_f.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
run_f.italic = True

p_solid = doc.add_paragraph()
p_solid.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_s = p_solid.add_run('Arquitectura bajo principios SOLID — Strategy Pattern — Multi-Bodega')
run_s.font.size = Pt(9)
run_s.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
run_s.italic = True

out = r'C:\Users\SSJUAN03\Desktop\WMS-PAME\Diseno_Unificado_Picking_Packing_v2.docx'
doc.save(out)
print(f'Documento generado: {out}')
