"""
Genera el documento de sesión WMS-PAME 2026-05-26.
Ejecutar: python generar_sesion_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Estilos globales ─────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1E, 0x83, 0x95)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x0B, 0x3D, 0x46)
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1E, 0x83, 0x95)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after  = Pt(2)
    p.add_run(text)
    return p

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x4A)
    return p

def divider(doc):
    doc.add_paragraph('─' * 80)

# ── PORTADA ──────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('WMS-PAME · Resumen de Sesión')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1E, 0x83, 0x95)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('26 de mayo de 2026')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x60, 0x80, 0x90)

doc.add_paragraph()

# Tabla resumen ejecutivo
tbl = doc.add_table(rows=4, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = [
    ('Repositorio',  'nsantgiral-dev/WMS-PAME  (rama: main)'),
    ('Rama',         'main'),
    ('Commits esta sesión', '4 commits (f2ae234, 2547836, 39b2745 + commit packing)'),
    ('Módulos tocados',     'despacho_parcial_service · pedidos_sync_service · packing (routes, mobile_service, app.js)'),
]
for i, (k, v) in enumerate(headers):
    row = tbl.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[0].paragraphs[0].runs[0].bold = True
    set_cell_bg(row.cells[0], 'D0EEF2')

doc.add_paragraph()

# ── SECCIÓN 1 ─────────────────────────────────────────────────────────────────
h1(doc, '1. Conector 244328 — Fix definitivo de f431_id_item')

h2(doc, '1.1 Contexto del problema')
body(doc,
    'El despacho parcial (flujo 244328 → 142945 → 142943) fallaba con HTTP 400 desde '
    'el conector 244328 con el error:')
code_block(doc, '"Compromisos: El item no existe por código, extensiones ni referencia"')
body(doc,
    'La causa raíz confirmada en Postman QA: el conector 244328 NO puede resolver ítems '
    'por f431_referencia_item (texto SKU como "PAPELSP9218"). Requiere obligatoriamente '
    'f431_id_item = ID numérico interno de Siesa (ej. 17368).')

h2(doc, '1.2 Causa raíz del commit 61e152d (el fix anterior que no funcionó)')
body(doc,
    'El commit 61e152d buscaba el ID numérico en PedidoSiesa.item_id_siesa (tabla local '
    'del read model). Había dos bugs que hacían que este campo estuviera NULL:')
bullet(doc, 'Bug A — El upsert del sync (pedidos_sync_service.py) NO actualizaba '
        'item_id_siesa en registros existentes. Solo se escribía en el INSERT inicial. '
        'Si el registro existía antes de que se agregara el campo, quedaba NULL para siempre.')
bullet(doc, 'Bug B — PedidoSiesa solo tiene ítems con cantidad_pendiente > 0. '
        'Si el pedido ya fue parcialmente despachado o si el sync no había corrido, '
        '_item_id_map quedaba vacío → fallback a referencia texto → mismo error 400.')

h2(doc, '1.3 Descubrimiento clave')
body(doc,
    'En Postman QA se confirmó que API_v2_Ventas_Pedidos_Compromisos retorna f120_id '
    '(ID numérico del ítem, campo nativo del JOIN T431 → T120) directamente en cada fila:')
code_block(doc, '{\n  "f431_rowid": 49387,\n  "f120_id": 11344,\n  "f120_referencia": "PAPELSP5765-",\n  "f405_cant_por_remisionar_base": 6.0000,\n  ...\n}')
body(doc,
    'compromisos_siesa ya estaba en memoria cuando se construía el payload de 244328. '
    'No era necesario consultar PedidoSiesa en absoluto.')

h2(doc, '1.4 Cambios aplicados — commit f2ae234')

h3(doc, 'despacho_parcial_service.py — _build_compromisos_244328()')
body(doc, 'Antes: buscaba id_item en item_id_map (PedidoSiesa.item_id_siesa, podía ser NULL).')
body(doc, 'Ahora: extrae f120_id directamente de la fila de compromisos_siesa.')
code_block(doc,
    '# Mapa {referencia: fila_completa} — f120_id nativo del JOIN T431→T120\n'
    '_comp_por_ref = {\n'
    '    str(r.get("f120_referencia", "")).strip(): r\n'
    '    for r in compromisos_siesa if r.get("f120_referencia")\n'
    '}\n\n'
    'for ref, rowid in rowid_map.items():\n'
    '    comp_row = _comp_por_ref.get(ref, {})\n'
    '    id_item  = comp_row.get("f120_id")  # ID numérico T120 — fuente primaria\n'
    '    cant_orig = float(comp_row.get("f405_cant_por_remisionar_base") or cant_real)\n'
    '    lote = comp_row.get("f405_id_lote") or None\n'
    '    ...'
)
body(doc,
    'Se eliminó el parámetro item_id_map y la query a PedidoSiesa. '
    'Se agregó warning de log si f120_id está ausente (diagnóstico temprano). '
    'También se propaga f405_id_lote desde compromisos.')

h3(doc, 'despacho_parcial_service.py — despachar_parcial()')
body(doc,
    'Se eliminaron las líneas que hacían la query a PedidoSiesa (_PS.query.filter_by...) '
    'y la construcción de _item_id_map. La llamada a _build_compromisos_244328 '
    'ya no necesita el parámetro item_id_map.')

h3(doc, 'pedidos_sync_service.py — upsert path')
body(doc, 'Se agregó la línea faltante en el path UPDATE del upsert:')
code_block(doc, 'reg.item_id_siesa = d["item_id_siesa"]   # f120_id — necesario para 244328')
body(doc,
    'Antes solo se escribía en el INSERT inicial. Ahora cada ciclo del scheduler '
    'actualiza item_id_siesa en los registros existentes, corrigiendo los NULL históricos.')

# ── SECCIÓN 2 ─────────────────────────────────────────────────────────────────
h1(doc, '2. Módulo Packing — Fix reiniciar-conteo')

h2(doc, '2.1 Problema reportado')
body(doc,
    'El botón "Reiniciar conteo" no funcionaba bien: reseteaba cantidad_real a 0 '
    'pero no actualizaba cantidad_esperada. El empacador veía las cantidades '
    'originales de Siesa, no las cantidades reales que el picker había recogido.')

h2(doc, '2.2 Causa raíz')
body(doc,
    'El endpoint PUT /api/packing/<id>/reiniciar-conteo hacía:')
bullet(doc, 'cantidad_real = 0 para todos los ítems ✅')
bullet(doc, 'verificado = False ✅')
bullet(doc, 'estado = EN_PROCESO ✅')
bullet(doc, 'cantidad_esperada: SIN TOCAR ❌ — quedaba con el valor del sync anterior o de Siesa')
body(doc, 'NO llamaba a PackingPickingSyncService.sincronizar() para re-sincronizar con picking.')

h2(doc, '2.3 Fix aplicado — commit 2547836')
body(doc, 'Después del reset del conteo (paso 1), se llama a sincronizar() (paso 2):')
code_block(doc,
    '# 1. Resetear conteo del empacador\n'
    'for item in tarea.items:\n'
    '    item.cantidad_real = 0\n'
    '    item.verificado    = False\n'
    'tarea.estado = EN_PROCESO\n'
    'db.session.commit()\n\n'
    '# 2. Re-sincronizar cantidad_esperada desde picking (estado ya EN_PROCESO)\n'
    'sync_resultado = PackingPickingSyncService.sincronizar(id)\n'
    '# silencioso si no hay picking asociado'
)
body(doc,
    'La respuesta del endpoint ahora incluye sync_picking con el resumen de cambios '
    '(items_ajustados, cambios con anterior/nuevo).')

# ── SECCIÓN 3 ─────────────────────────────────────────────────────────────────
h1(doc, '3. Módulo Packing — Fix HUD Empacador (nombres y conteo)')

h2(doc, '3.1 Problema reportado')
body(doc,
    'Con 2 referencias distintas en el packing (ej. 2 pq marcador Pelikan + 2 resmas carta):')
bullet(doc, 'Los nombres de productos se "bugueaban" — el nombre del producto anterior '
        'aparecía brevemente después de escanear.')
bullet(doc, 'El conteo no era consistente: escaneando 1 marcador cuando se esperaban 2, '
        'el HUD saltaba prematuramente a la resma dejando 1 marcador sin contar.')

h2(doc, '3.2 Bug A — empFlash en app.js (nombres que revierten)')
body(doc, 'Secuencia del bug con scan de empaque (DUN-14/LPN):')
bullet(doc, 'empFlash() captura prevText = "MARCADOR PELIKAN" (nombre actual en DOM)', 0)
bullet(doc, 'msgEl.textContent = "PIEZA completa — 12 und" (flash)', 0)
bullet(doc, 'empRenderHUDItem() avanza al siguiente ítem → DOM = "RESMA CARTA"', 0)
bullet(doc, '900ms después: setTimeout → msgEl.textContent = prevText = "MARCADOR PELIKAN" ← BUG', 0)
body(doc, 'Fix: en el timeout verde, reemplazar msgEl.textContent = prevText por empRenderHUDItem():')
code_block(doc,
    '// ANTES\n'
    'setTimeout(() => {\n'
    '  msgEl.style.color = "";\n'
    '  msgEl.textContent = prevText;  // ← nombre incorrecto\n'
    '}, 900);\n\n'
    '// AHORA\n'
    'setTimeout(() => {\n'
    '  msgEl.style.color = "";\n'
    '  empRenderHUDItem();  // ← re-renderiza ítem actual correcto\n'
    '}, 900);'
)

h2(doc, '3.3 Bug B — item.verificado en mobile_service.py (conteo prematuro)')
body(doc, 'Antes: item.verificado = item.cantidad_real is not None')
body(doc,
    'is not None es True después del PRIMER scan, aunque cantidad_real < cantidad_esperada. '
    'Ejemplo: esperado=2 marcadores, packer escanea 1 → verificado=True → HUD salta '
    'a RESMA CARTA → segundo marcador nunca se cuenta.')
body(doc, 'Fix: item.verificado = nueva >= item.cantidad_esperada')
code_block(doc,
    '# ANTES\n'
    'item.cantidad_real = nueva\n'
    'item.verificado = item.cantidad_real is not None  # True tras 1er scan\n\n'
    '# AHORA\n'
    'item.cantidad_real = nueva\n'
    'item.verificado = nueva >= item.cantidad_esperada  # True solo al alcanzar el esperado'
)

h2(doc, '3.4 Comportamiento nuevo')
tbl2 = doc.add_table(rows=4, cols=3)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
encabezados = ['Situación', 'Antes', 'Ahora']
for i, h in enumerate(encabezados):
    cell = tbl2.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    set_cell_bg(cell, 'D0EEF2')
filas = [
    ('Scan 1 de 2 marcadores',     'HUD salta a RESMA (verificado=True)', 'HUD queda en MARCADOR, muestra 1 de 2 ✅'),
    ('Scan 2 de 2 marcadores',     '—',                                     'HUD avanza a RESMA ✅'),
    ('Empaque DUN-14 (24 und)',     'Nombre revierte 900ms después',         'Nombre correcto al desaparecer el flash ✅'),
]
for i, (s, a, n) in enumerate(filas, start=1):
    row = tbl2.rows[i]
    row.cells[0].text = s
    row.cells[1].text = a
    row.cells[2].text = n

doc.add_paragraph()

# ── SECCIÓN 4 — Commits ───────────────────────────────────────────────────────
h1(doc, '4. Tabla de commits de la sesión')

tbl3 = doc.add_table(rows=5, cols=3)
tbl3.style = 'Table Grid'
for i, h in enumerate(['Commit', 'Archivo(s)', 'Descripción']):
    cell = tbl3.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    set_cell_bg(cell, 'D0EEF2')
commits = [
    ('f2ae234', 'despacho_parcial_service.py\npedidos_sync_service.py',
     'fix(244328): usar f120_id de compromisos_siesa en vez de PedidoSiesa'),
    ('2547836', 'app/routes/packing.py',
     'fix(packing): reiniciar-conteo re-sincroniza cantidad_esperada desde picking'),
    ('39b2745', 'app/static/pwa/app.js\napp/services/mobile_service.py',
     'fix(packing): corregir bug nombres y conteo incremental en HUD empacador'),
    ('(anterior)', 'app/routes/packing.py',
     'fix(packing): reiniciar-conteo actualizado con sync picking'),
]
for i, (c, f, d) in enumerate(commits, start=1):
    row = tbl3.rows[i]
    run = row.cells[0].paragraphs[0].add_run(c)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    row.cells[1].text = f
    row.cells[2].text = d

doc.add_paragraph()

# ── SECCIÓN 5 — Próximos pasos ────────────────────────────────────────────────
h1(doc, '5. Próximos pasos / Pendientes')

bullet(doc, 'Verificar en Railway que el deploy de 39b2745 esté activo (~2-3 min tras push)')
bullet(doc, 'Prueba end-to-end despacho parcial: POST /api/despacho_parcial/<id>/despachar '
        '→ verificar que 244328 retorna "Transacción Exitosa" en logs de Railway')
bullet(doc, 'Verificar en Siesa que f405_cant_por_remisionar_base en T405 = cantidad real picada '
        '(no la comprometida original)')
bullet(doc, 'Verificar en Siesa que T470 registra la cantidad correcta (3, no 4 en el caso de prueba)')
bullet(doc, 'Una vez QA pasa: replicar en producción (servicios.siesacloud.com)')
bullet(doc, 'Desactivar consulta 7811 en Connekta (ya no se necesita — reemplazada por 244328)')
bullet(doc, 'Prueba packing con 2 referencias distintas: escanear marcadores (ambas unidades) '
        '→ verificar que el HUD no salta prematuramente ni invierte nombres')

# ── PIE ───────────────────────────────────────────────────────────────────────
doc.add_paragraph()
divider(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Generado automáticamente · WMS-PAME · {datetime.date.today().strftime("%d/%m/%Y")}')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x80, 0x9A, 0xAA)

# ── Guardar ───────────────────────────────────────────────────────────────────
out = r'C:\Users\SSJUAN03\Desktop\WMS-PAME\Sesion_WMS_PAME_2026-05-26.docx'
doc.save(out)
print(f'Documento generado: {out}')
