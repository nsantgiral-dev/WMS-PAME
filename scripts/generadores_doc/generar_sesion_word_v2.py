"""Genera el documento de sesión WMS-PAME 2026-06-03."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)
    return p

def add_table(doc, headers, rows, col_widths=None, hdr='1F497D'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr_cells[i], hdr)
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        rc = table.add_row().cells
        for i, val in enumerate(row_data):
            rc[i].paragraphs[0].add_run(val).font.size = Pt(9)
    if col_widths:
        from docx.oxml import OxmlElement as OE
        from docx.oxml.ns import qn as QN
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcW = OE('w:tcW')
                    tcW.set(QN('w:w'), str(int(col_widths[j] * 567)))
                    tcW.set(QN('w:type'), 'dxa')
                    tcPr.append(tcW)
    return table

# ══ PORTADA ══════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(50)
r = p.add_run('WMS-PAME')
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = RGBColor(0x1F,0x49,0x7D)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Resumen de Sesión — 03 de Junio 2026')
r2.bold = True; r2.font.size = Pt(16); r2.font.color.rgb = RGBColor(0x2E,0x74,0xB5)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Unificación Picking/Packing PD + ST · Flujo de Traslados · Módulo Tienda')
r3.font.size = Pt(12); r3.font.color.rgb = RGBColor(0x60,0x60,0x60)

doc.add_page_break()

# ══ 1. CONTEXTO ══════════════════════════════════════════════════════════════
h1(doc, '1. Contexto de la Sesión')
body(doc, 'Esta sesión partió del análisis exhaustivo de dos flujos separados que existían '
          'en el WMS-PAME: el flujo de Pedidos (PD) con picking → packing → cierre → '
          'Factura+RM en SIESA, y el flujo de Requisiciones/Traslados (ST) con su propio '
          'picking y packing aislados. El objetivo fue unificarlos en una sola pantalla '
          'para el operario, respetando los principios SOLID y sin romper nada existente.')

# ══ 2. LO QUE SE ANALIZÓ ════════════════════════════════════════════════════
h1(doc, '2. Análisis Realizado')

h2(doc, '2.1  Auditoría de los dos flujos')
body(doc, 'Se auditaron exhaustivamente los dos flujos antes de escribir código:')
bullet(doc, 'Flujo PD: sync SIESA → PedidoSiesa → picking ciego → packing doble conteo → '
           'cerrar caja → SiesaJob DESPACHO_F470 → 238925 (Factura FE + Remisión RM)')
bullet(doc, 'Flujo ST: tienda crea solicitud → admin aprueba → picking → '
           'confirmar_picking_traslado → 174646 (RIT) → packing → 174720 (Compromisos) → '
           'despachar → 174930/173076 (STS) → tienda destino → 173079 (ETS)')
bullet(doc, 'Hallazgo crítico: los traslados NO tenían packing real — confirmar_packing era '
           'un bypass directo a SIESA sin doble conteo del empacador.')
bullet(doc, 'La RIT (174646) se creaba después del picking, lo que dejaba el inventario '
           'sin reservar durante el período de recogida.')

h2(doc, '2.2  APIs SIESA estudiadas')
add_table(doc,
    ['Conector', 'Nombre', 'Rol en el flujo'],
    [
        ['174646', 'RequisicionesParaTransferir (RIT)', 'Reserva inventario en SIESA al aprobar'],
        ['174720', 'CompromisosDesdeRequisicion', 'Cantidades reales + ubicaciones (al cerrar packing)'],
        ['174930', 'TransferenciasDesdeRequisicion (STS)', 'Crea salida en tránsito desde RIT'],
        ['173076', 'TransferenciaEnTransitoSalida', 'Fallback: salida manual sin RIT'],
        ['173079', 'TransferenciaEnTransitoEntrada (ETS)', 'Tienda destino confirma recepción'],
        ['176',    'GET Transferencia_Salida_Transito', 'Consulta: verificar STS antes de recibir'],
        ['177',    'GET Transferencia_Transito_Entrada', 'Consulta: verificar ETS creada'],
    ],
    col_widths=[2, 6, 8.5]
)

h2(doc, '2.3  Módulo de Usuarios — diagnóstico')
body(doc, 'Se revisó el formulario de creación de usuarios y se detectaron dos gaps:')
bullet(doc, 'Los roles picker_traslado y packer_traslado existían en el backend pero apuntaban '
           'a pantallas separadas que serían eliminadas por el nuevo diseño unificado.')
bullet(doc, 'Los nombres de bodegas en _BODEGAS_ORIGEN (frontend) no coincidían con los del '
           'formulario de usuario (NS1 era "Neiva Sur" vs "Neiva Sur Principal").')

# ══ 3. DECISIONES ARQUITECTÓNICAS ════════════════════════════════════════════
h1(doc, '3. Decisiones Arquitectónicas Tomadas')

add_table(doc,
    ['Decisión', 'Justificación SOLID'],
    [
        ['Strategy Pattern para cerrar_packing(): PedidoPackingCloser + TrasladoPackingCloser',
         'O — Open/Closed: agregar tipo futuro = nuevo archivo, sin tocar packing_service'],
        ['IPackingCloser como ABC + PackingCloserFactory',
         'D — Dependency Inversion: packing_service depende de la abstracción, no de implementaciones'],
        ['scoping/task_scope.py: scope_picking() + scope_packing()',
         'S — Single Responsibility: un solo lugar decide qué ve cada usuario'],
        ['RIT (174646) movida a aprobar_solicitud (no a confirmar_picking)',
         'El inventario debe bloquearse en SIESA ANTES de que el picker empiece, no después'],
        ['picker_traslado y packer_traslado → pantallas unificadas',
         'Elimina duplicación; los roles mantienen significado como operarios de tienda'],
        ['bodega_origen_siesa como campo de scoping en TareaPicking/TareaPacking',
         'Permite N tiendas sin cambios de código — scoping genérico por valor de campo'],
    ],
    col_widths=[8, 8.5]
)

# ══ 4. CAMBIOS IMPLEMENTADOS ═════════════════════════════════════════════════
h1(doc, '4. Cambios Implementados — Commits')

add_table(doc,
    ['Commit', 'Descripción', 'Archivos clave'],
    [
        ['321b969', 'feat(tienda): "Pedir desde" + inversión origen/destino\nSelector apunta a la bodega fuente; tienda logueada = destino automático',
         'app.js, index.html, traslados.py, traslado_service.py'],
        ['206f8c5', 'fix(tienda): alinear nombres bodegas _BODEGAS_ORIGEN con formulario usuario\nNS1 = "Neiva Sur Principal"; orden FC1 antes que PC1',
         'app.js'],
        ['9bfc4c3', 'fix(traslados): mover RIT (174646) de picking → aprobación\nInventario reservado en SIESA antes de que el picker empiece',
         'traslado_service.py'],
        ['e8cd367', 'feat(migrations): picking/packing unificado — campos ST\nMigración b3c4d5e6f7g8: bodega_origen_siesa + tipo_documento + referencia_doc + solicitud_id + tienda_destino',
         'migrations/versions/b3c4d5e6f7g8_*.py'],
        ['8903314', 'feat(packing): Strategy Pattern PD + ST\nclosing/ con IPackingCloser, PedidoPackingCloser, TrasladoPackingCloser, factory\nscoping/task_scope.py\npacking_service delega al factory',
         'closing/*.py, scoping/task_scope.py, packing_service.py, packing.py, picking.py'],
        ['549eb05', 'feat(picking-packing): scoping, lock 409, TareaPacking ST, DLQ TRASLADO\nPickingService con bodega_origen_siesa; confirmar_picking crea TareaPacking; DLQ handler 174930',
         'picking_service.py, traslado_service.py, siesa_job_service.py, routes/picking.py, routes/packing.py'],
        ['af25836', 'feat(frontend): etiquetas [PEDIDO]/[TRASLADO] + prioridad PD>ST\nPacking list: etiqueta + tienda destino para ST; pantalla operario: etiqueta en tarea; cola SQL PD primero',
         'app.js, mobile_service.py'],
        ['eb4d4cb', 'refactor(roles): picker_traslado y packer_traslado → pantallas unificadas\nElimina pantallas separadas; roles redefinidos como operarios de tienda con scoping automático ST',
         'index.html, app.js, task_scope.py, mobile_service.py'],
    ],
    col_widths=[1.8, 7.2, 7.5]
)

# ══ 5. ESTADO FINAL DEL FLUJO ═════════════════════════════════════════════════
h1(doc, '5. Estado Final del Flujo Unificado')

h2(doc, '5.1  Flujo Pedido (PD) — sin cambios funcionales')
code_block(doc,
    'SIESA sync → PedidoSiesa\n'
    '  Admin: iniciar-despacho → TareaPicking[] (tipo=PEDIDO, bodega=NB1)\n'
    '  Operario NB1: cola muestra [PEDIDO] primero (prioridad 1)\n'
    '    Lock 409 si tiene tarea EN_PROCESO\n'
    '  Admin: crear-desde-picking → TareaPacking (tipo=PEDIDO, referencia=PD1307)\n'
    '  Empacador NB1: cola unificada — etiqueta azul [PEDIDO]\n'
    '    cerrar_packing() → factory.get("PEDIDO") → PedidoPackingCloser\n'
    '    → SiesaJob DESPACHO_F470 → DLQ → 238925 → Factura(FE) + RM'
)

h2(doc, '5.2  Flujo Requisición / Traslado (ST) — rediseñado')
code_block(doc,
    'Tienda (cualquier bodega): "Pedir desde" → selecciona origen → carrito\n'
    '  POST /api/traslados/ → SolicitudTraslado BORRADOR → ENVIADA\n'
    '  Admin aprueba → TareaPicking[] (tipo=TRASLADO, bodega=origen)\n'
    '               → 174646 (RIT) disparada → inventario SIESA = Comprometido\n'
    '  Operario (NB1 o tienda según origen):\n'
    '    cola muestra [TRASLADO] naranja (después de PD si coexisten)\n'
    '    conteo ciego → confirmar picking\n'
    '    → TareaPacking (tipo=TRASLADO, solicitud_id, tienda_destino) creada\n'
    '  Empacador: cola unificada — etiqueta naranja [TRASLADO] + "→ Tienda X"\n'
    '    cerrar_packing() → factory.get("TRASLADO") → TrasladoPackingCloser\n'
    '    → 174720 (Compromisos sync) + SiesaJob DESPACHO_TRASLADO → DLQ → 174930 (STS)\n'
    '    → SolicitudTraslado: EN_TRANSITO\n'
    '  Tienda destino: tab "Recibir" → confirmar_recepcion\n'
    '    → 173079 (ETS) → ENTREGADA → RIT Estado 4 (Cumplido)'
)

h2(doc, '5.3  Scoping por rol — quién ve qué')
add_table(doc,
    ['Rol', 'Picking ve', 'Packing ve', 'PD', 'Conteo'],
    [
        ['admin / supervisor', 'Todo', 'Todo', 'Sí', 'Sí'],
        ['operario (NB1)', 'PD + ST bodega NB1', 'No', 'Sí', 'Sí'],
        ['empacador (NB1)', 'No', 'PD + ST bodega NB1', 'Sí', 'No'],
        ['picker_traslado (tienda)', 'Solo ST bodega_siesa_id', 'No', 'No', 'No'],
        ['packer_traslado (tienda)', 'No', 'Solo ST bodega_siesa_id', 'No', 'No'],
        ['tienda (PV)', 'Solo ST origen de su bodega', 'Solo ST origen de su bodega', 'No', 'No'],
    ],
    col_widths=[4, 3.5, 3.5, 1.2, 1.8]
)

# ══ 6. ARQUITECTURA DE ARCHIVOS NUEVOS ════════════════════════════════════════
h1(doc, '6. Arquitectura de Archivos Creados')

add_table(doc,
    ['Archivo', 'Tipo', 'Propósito'],
    [
        ['app/services/closing/__init__.py', 'Nuevo', 'Módulo closing'],
        ['app/services/closing/base.py', 'Nuevo', 'IPackingCloser (ABC) + CierreResult dataclass'],
        ['app/services/closing/pedido_closer.py', 'Nuevo', 'PedidoPackingCloser → 238925 DLQ'],
        ['app/services/closing/traslado_closer.py', 'Nuevo', 'TrasladoPackingCloser → 174720 sync + DESPACHO_TRASLADO DLQ'],
        ['app/services/closing/factory.py', 'Nuevo', 'PackingCloserFactory.get(tipo_documento)'],
        ['app/services/scoping/__init__.py', 'Nuevo', 'Módulo scoping'],
        ['app/services/scoping/task_scope.py', 'Nuevo', 'scope_picking() + scope_packing() centralizados'],
        ['migrations/versions/b3c4d5e6f7g8_*.py', 'Nuevo', 'Migración: bodega_origen_siesa + campos ST en packing'],
    ],
    col_widths=[6.5, 1.8, 8.2]
)

# ══ 7. PENDIENTES PARA PRÓXIMA SESIÓN ════════════════════════════════════════
h1(doc, '7. Pendientes para la Próxima Sesión')

h2(doc, '7.1  Verificación de la migración en Railway')
body(doc, 'La migración b3c4d5e6f7g8 debe ejecutarse en Railway en el próximo deploy. '
          'Verificar en los logs de Railway que "flask db upgrade" completó sin errores '
          'antes de continuar con cualquier prueba de funcionalidad.')
bullet(doc, 'Confirmar que tareas_picking tiene columna bodega_origen_siesa')
bullet(doc, 'Confirmar que tareas_packing tiene las 5 columnas nuevas')
bullet(doc, 'Confirmar que numero_pedido_siesa es nullable en producción')

h2(doc, '7.2  Formulario de Creación de Usuarios — gaps detectados')
body(doc, 'El formulario de creación de usuarios tiene dos problemas no resueltos:')
bullet(doc, 'GAP 1: El campo bodega_siesa_id no se muestra ni se puede asignar cuando el rol '
           'es picker_traslado o packer_traslado (solo aparece para rol=tienda). '
           'Estos roles también necesitan bodega_siesa_id para que el scoping funcione.')
bullet(doc, 'GAP 2: El campo Conteos cíclicos por día aparece para todos los roles incluyendo '
           'tiendas, donde no tiene sentido. Debería ocultarse para roles de tienda.')

h2(doc, '7.3  Tab "Recibir" en el módulo de tienda')
body(doc, 'Verificar que el tab "Recibir" en la pantalla de tienda sigue funcionando '
          'correctamente después de los cambios. Este tab llama a confirmar_recepcion() '
          'que dispara el 173079 (ETS). No se modificó pero debe validarse.')

h2(doc, '7.4  Panel admin — requisiciones con TareaPacking')
body(doc, 'El panel de admin de requisiciones (tab Requisiciones) muestra el estado de las '
          'solicitudes de traslado. Con el nuevo flujo, una solicitud EN_PACKING tiene una '
          'TareaPacking asociada. Verificar que el panel refleja correctamente este estado '
          'y que el admin puede ver qué tareas de packing están asociadas.')

h2(doc, '7.5  Prueba end-to-end del flujo ST unificado')
body(doc, 'El flujo completo no ha sido probado en producción. El orden de validación sugerido:')
add_table(doc,
    ['#', 'Paso a validar', 'Qué verificar'],
    [
        ['1', 'Tienda crea solicitud con "Pedir desde" NB1', 'bodega_origen = NB1, destino = tienda logueada'],
        ['2', 'Admin aprueba', 'RIT 174646 disparada; TareaPicking con bodega_origen_siesa=NB1'],
        ['3', 'Operario NB1 hace picking', 'Ve etiqueta [TRASLADO]; lock 409 si tiene tarea activa'],
        ['4', 'Confirmar picking', 'TareaPacking creada (tipo=TRASLADO); solicitud → EN_PACKING'],
        ['5', 'Empacador NB1 hace packing', 'Ve etiqueta naranja [TRASLADO] + tienda destino'],
        ['6', 'Cerrar caja', '174720 sync + SiesaJob DESPACHO_TRASLADO encolado'],
        ['7', 'DLQ procesa DESPACHO_TRASLADO', '174930 disparado; solicitud → EN_TRANSITO'],
        ['8', 'Tienda recibe', '173079 (ETS) disparado; solicitud → ENTREGADA'],
    ],
    col_widths=[0.8, 5, 10.7]
)

h2(doc, '7.6  Módulo de Packing Admin — lista unificada')
body(doc, 'El panel admin de packing (tab Bodega) actualmente filtra tareas manualmente '
          'pasando almacen_id desde el frontend. Con el nuevo scoping server-side, '
          'este filtro cliente ya no es necesario. Revisar si el admin ve correctamente '
          'tanto PD como ST en su lista de packing.')

h2(doc, '7.7  Deuda técnica generada')
add_table(doc,
    ['Deuda', 'Archivo', 'Descripción'],
    [
        ['DT-PACK-01', 'packing_service.py',
         '_cerrar_packing_pedido_legacy() quedó como método huérfano. Eliminar cuando se confirme que PedidoPackingCloser funciona en producción.'],
        ['DT-PACK-02', 'traslados.py',
         'Endpoints /cola-picker y /cola-packer siguen existiendo pero ya no los usa ningún frontend. Pueden eliminarse cuando se confirme que picker/packer_traslado funcionan en pantallas unificadas.'],
        ['DT-PACK-03', 'traslado_service.py',
         'confirmar_packing_traslado() sigue existiendo para compatibilidad con llamadas admin directas. Revisar si puede eliminarse o simplificarse.'],
        ['DT-PACK-04', 'app.js',
         'TRAS_PICK = null y funciones trasPickerCargarCola / cargarTrasladosOperario siguen referenciadas en el código por cargarTrasladosOperario(). Limpiar en próxima sesión.'],
    ],
    col_widths=[2.5, 4, 10]
)

# ══ 8. ARQUITECTURA RESUMEN ════════════════════════════════════════════════════
h1(doc, '8. Arquitectura Resultante — Diagrama de Cierre')

code_block(doc,
    'CERRAR CAJA (mismo botón para todos)\n'
    '           ↓\n'
    'PackingService.cerrar_packing(tarea_id, bultos, usuario_id)\n'
    '           ↓\n'
    'PackingCloserFactory.get(tarea.tipo_documento)\n'
    '        ↙                        ↘\n'
    '   PEDIDO                     TRASLADO\n'
    'PedidoPackingCloser      TrasladoPackingCloser\n'
    '     ↓                         ↓ sync     ↓ DLQ\n'
    'SiesaJob                     174720     SiesaJob\n'
    'DESPACHO_F470            Compromisos  DESPACHO_TRASLADO\n'
    '     ↓                                    ↓\n'
    '  238925                               174930\n'
    'Factura(FE)                         STS (Salida)\n'
    'Remisión(RM)                        EN_TRANSITO\n'
    '     ↓                                    ↓\n'
    'Cliente recibe              Tienda destino confirma\n'
    '                                     ↓\n'
    '                               173079 (ETS)\n'
    '                               ENTREGADA\n'
    '                           RIT → Estado 4 Cumplido'
)

# ══ PIE ═══════════════════════════════════════════════════════════════════════
doc.add_page_break()
p_f = doc.add_paragraph()
p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_f.paragraph_format.space_before = Pt(60)
rf = p_f.add_run('WMS-PAME  |  Sesión 2026-06-03  |  Principios SOLID · Strategy Pattern · Multi-Bodega')
rf.font.size = Pt(9)
rf.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
rf.italic = True

out = r'C:\Users\SSJUAN03\Desktop\WMS-PAME\Sesion_WMS_PAME_2026-06-03.docx'
doc.save(out)
print(f'Generado: {out}')
